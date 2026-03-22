# ============================================================
# MERLIN RAG API — Enchanted Forest Edition
# FastAPI + Qdrant + Ollama with auth & conversation history
# ============================================================

import os
import re
import json
import time
import asyncio
import sqlite3
import shutil
import uuid
import hashlib
import hmac
import csv
import io
import base64
import tempfile
from datetime import datetime
from typing import List, Optional, Dict, Any, Tuple
from urllib.parse import urlparse, unquote

import requests
import bcrypt
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
from fastapi import FastAPI, HTTPException, UploadFile, File, Request, Response, Depends, BackgroundTasks, Form
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pathlib import Path
from pydantic import BaseModel
from qdrant_client import QdrantClient
from qdrant_client.http import models as qm
from crawl4ai import AsyncWebCrawler, CacheMode, CrawlerRunConfig
from crawl4ai.content_filter_strategy import PruningContentFilter
from crawl4ai.markdown_generation_strategy import DefaultMarkdownGenerator
from docx import Document as DocxDocument
from openpyxl import load_workbook
import pytesseract
from PIL import Image
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

# ---------- Configuration ----------
APP_TITLE = "Merlin RAG API"
DATA_DIR = os.getenv("RAG_DATA_DIR", "/rag-data/text")
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://ollama-llm:11434")
CHAT_MODEL = os.getenv("CHAT_MODEL", "llama3.1:8b")
CHAT_API_URL = os.getenv("CHAT_API_URL", "")  # vLLM OpenAI-compatible endpoint; empty = use Ollama
CHAT_API_KEY = os.getenv("CHAT_API_KEY", "")
EMBED_MODEL = os.getenv("EMBED_MODEL", "nomic-embed-text")
QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant:6333")
COLLECTION = os.getenv("QDRANT_COLLECTION", "general_docs")
SQLITE_PATH = os.getenv("SQLITE_PATH", "/data/chat.db")
RAG_ROOT = Path(os.getenv("RAG_ROOT", "/rag-data"))
TEXT_DIR = Path(DATA_DIR)
INCOMING_DIR = RAG_ROOT / "incoming"
TEXT_DIR.mkdir(parents=True, exist_ok=True)
INCOMING_DIR.mkdir(parents=True, exist_ok=True)
PROJECTS_DIR = RAG_ROOT / "projects"
PROJECTS_DIR.mkdir(parents=True, exist_ok=True)
MIN_CONTEXT_CHARS = int(os.getenv("MIN_CONTEXT_CHARS", "1200"))
MIN_TOP_SCORE = float(os.getenv("MIN_TOP_SCORE", "0.25"))

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "900"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "150"))
TOP_K = int(os.getenv("TOP_K", "40"))
#adjust the below based on what your HW can handle.  i'm running mine a bit sporty
WEB_SEARCH_RESULTS_DEFAULT = int(os.getenv("WEB_SEARCH_RESULTS_DEFAULT", "5"))
WEB_SEARCH_RESULTS_MAX = int(os.getenv("WEB_SEARCH_RESULTS_MAX", "8"))
MODEL_MAX_INPUT_TOKENS = int(os.getenv("MODEL_MAX_INPUT_TOKENS", "16092"))
MODEL_OUTPUT_TOKEN_RESERVE = int(os.getenv("MODEL_OUTPUT_TOKEN_RESERVE", "10024"))
APPROX_CHARS_PER_TOKEN = max(1, int(os.getenv("APPROX_CHARS_PER_TOKEN", "4")))
DOC_CONTEXT_TOKEN_BUDGET = int(os.getenv("DOC_CONTEXT_TOKEN_BUDGET", "16000"))
WEB_CONTEXT_TOKEN_BUDGET = int(os.getenv("WEB_CONTEXT_TOKEN_BUDGET", "12000"))
HISTORY_TOKEN_BUDGET = int(os.getenv("HISTORY_TOKEN_BUDGET", "1200"))
# BRAVE_SEARCH_API_KEY = os.getenv("BRAVE_SEARCH_API_KEY", "")
# BRAVE_SEARCH_BASE_URL = os.getenv("BRAVE_SEARCH_BASE_URL", "https://api.search.brave.com/res/v1/web/search")
# BRAVE_SAFESEARCH = os.getenv("BRAVE_SAFESEARCH", "moderate")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "")
TAVILY_SEARCH_BASE_URL = os.getenv("TAVILY_SEARCH_BASE_URL", "https://api.tavily.com/search")
TAVILY_SEARCH_TOPIC = os.getenv("TAVILY_SEARCH_TOPIC", "general")
WEB_SEARCH_COUNTRY = os.getenv("WEB_SEARCH_COUNTRY", "US")
WEB_SEARCH_UI_LANG = os.getenv("WEB_SEARCH_UI_LANG", "en-US")
WEB_SEARCH_PAGE_CHAR_LIMIT = int(os.getenv("WEB_SEARCH_PAGE_CHAR_LIMIT", "4000"))
WEB_SNIPPET_CHAR_LIMIT = int(os.getenv("WEB_SNIPPET_CHAR_LIMIT", "600"))

VISION_MODEL = os.getenv("VISION_MODEL", "llava:7b")
OCR_MIN_CHARS = int(os.getenv("OCR_MIN_CHARS", "50"))
EXPORTS_DIR = Path(os.getenv("EXPORTS_DIR", "/tmp/merlin-exports"))
EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
EXPORT_MAX_AGE = int(os.getenv("EXPORT_MAX_AGE", "3600"))

TEXT_EXTENSIONS = (".md", ".txt", ".docx", ".xlsx", ".csv", ".pdf")
IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff")

# Auth configuration
SESSION_SECRET = os.getenv(
    "SESSION_SECRET",
    hashlib.sha256(f"merlin-dev-{os.getenv('HOSTNAME', 'local')}".encode()).hexdigest(),
)
SESSION_MAX_AGE = int(os.getenv("SESSION_MAX_AGE", "28800"))  # 8 hours
INTERNAL_KEY = os.getenv("INTERNAL_API_KEY", "")
COOKIE_SECURE = os.getenv("COOKIE_SECURE", "false").lower() == "true"
serializer = URLSafeTimedSerializer(SESSION_SECRET)

# ---------- FastAPI Application ----------
app = FastAPI(title=APP_TITLE)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://127.0.0.1:8000", "http://localhost:8000",
        "http://127.0.0.1:3000", "http://localhost:3000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- Database ----------
def db() -> sqlite3.Connection:
    os.makedirs(os.path.dirname(SQLITE_PATH), exist_ok=True)
    conn = sqlite3.connect(SQLITE_PATH)
    conn.execute("PRAGMA foreign_keys = ON")

    conn.execute("""
      CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL,
        created_at INTEGER NOT NULL
      )
    """)

    conn.execute("""
      CREATE TABLE IF NOT EXISTS projects (
        id TEXT PRIMARY KEY,
        user_id INTEGER NOT NULL,
        name TEXT NOT NULL,
        created_at INTEGER NOT NULL,
        updated_at INTEGER NOT NULL,
        FOREIGN KEY (user_id) REFERENCES users(id)
      )
    """)
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_projects_user "
        "ON projects(user_id, updated_at DESC)"
    )

    conn.execute("""
      CREATE TABLE IF NOT EXISTS project_files (
        id TEXT PRIMARY KEY,
        project_id TEXT NOT NULL,
        user_id INTEGER NOT NULL,
        filename TEXT NOT NULL,
        stored_name TEXT NOT NULL,
        rel_path TEXT NOT NULL,
        media_kind TEXT NOT NULL,
        size INTEGER NOT NULL,
        created_at INTEGER NOT NULL,
        updated_at INTEGER NOT NULL,
        FOREIGN KEY (project_id) REFERENCES projects(id) ON DELETE CASCADE,
        FOREIGN KEY (user_id) REFERENCES users(id)
      )
    """)
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_project_files_project "
        "ON project_files(project_id, updated_at DESC)"
    )

    conn.execute("""
      CREATE TABLE IF NOT EXISTS conversations (
        id TEXT PRIMARY KEY,
        user_id INTEGER NOT NULL,
        project_id TEXT,
        title TEXT NOT NULL DEFAULT 'New Conversation',
        created_at INTEGER NOT NULL,
        updated_at INTEGER NOT NULL,
        FOREIGN KEY (user_id) REFERENCES users(id),
        FOREIGN KEY (project_id) REFERENCES projects(id)
      )
    """)
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_conversations_user "
        "ON conversations(user_id, updated_at DESC)"
    )

    conn.execute("""
      CREATE TABLE IF NOT EXISTS messages (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        convo_id TEXT NOT NULL,
        role TEXT NOT NULL,
        content TEXT NOT NULL,
        ts INTEGER NOT NULL
      )
    """)
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_messages_convo_ts ON messages(convo_id, ts)"
    )

    # Migrate: add user_id column if missing
    cursor = conn.execute("PRAGMA table_info(messages)")
    columns = [row[1] for row in cursor.fetchall()]
    if "user_id" not in columns:
        conn.execute("ALTER TABLE messages ADD COLUMN user_id INTEGER")

    cursor = conn.execute("PRAGMA table_info(conversations)")
    convo_columns = [row[1] for row in cursor.fetchall()]
    if "project_id" not in convo_columns:
        conn.execute("ALTER TABLE conversations ADD COLUMN project_id TEXT")

    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_conversations_project "
        "ON conversations(user_id, project_id, updated_at DESC)"
    )

    return conn


def add_message(convo_id: str, role: str, content: str, user_id: int = None) -> None:
    conn = db()
    with conn:
        conn.execute(
            "INSERT INTO messages (convo_id, role, content, ts, user_id) VALUES (?, ?, ?, ?, ?)",
            (convo_id, role, content, int(time.time()), user_id),
        )
    conn.close()


def get_recent_messages(convo_id: str, limit: int = 12) -> List[Dict[str, str]]:
    conn = db()
    cur = conn.execute(
        "SELECT role, content FROM messages WHERE convo_id=? ORDER BY ts DESC, id DESC LIMIT ?",
        (convo_id, limit),
    )
    rows = cur.fetchall()
    conn.close()
    return [{"role": r[0], "content": r[1]} for r in reversed(rows)]


def user_project_root(user_id: int, project_id: str) -> Path:
    return PROJECTS_DIR / str(user_id) / project_id


def project_text_dir(user_id: int, project_id: str) -> Path:
    path = user_project_root(user_id, project_id) / "text"
    path.mkdir(parents=True, exist_ok=True)
    return path


def project_incoming_dir(user_id: int, project_id: str) -> Path:
    path = user_project_root(user_id, project_id) / "incoming"
    path.mkdir(parents=True, exist_ok=True)
    return path


def get_project_row(conn: sqlite3.Connection, project_id: str, user_id: int):
    cur = conn.execute(
        "SELECT id, user_id, name, created_at, updated_at FROM projects WHERE id=? AND user_id=?",
        (project_id, user_id),
    )
    return cur.fetchone()


def require_project(conn: sqlite3.Connection, project_id: str, user_id: int):
    row = get_project_row(conn, project_id, user_id)
    if not row:
        raise HTTPException(status_code=404, detail="Project not found")
    return row


# ---------- Authentication ----------
def create_session_cookie(user_id: int, username: str) -> str:
    return serializer.dumps({"user_id": user_id, "username": username})


def verify_session_cookie(cookie_value: str) -> Optional[Dict]:
    try:
        return serializer.loads(cookie_value, max_age=SESSION_MAX_AGE)
    except (BadSignature, SignatureExpired):
        return None


def get_current_user(request: Request) -> Dict:
    cookie = request.cookies.get("merlin_session")
    if cookie:
        user = verify_session_cookie(cookie)
        if user:
            return user
    raise HTTPException(status_code=401, detail="Not authenticated")


def get_optional_user(request: Request) -> Dict:
    internal = request.headers.get("X-Internal-Key", "")
    if INTERNAL_KEY and hmac.compare_digest(internal, INTERNAL_KEY):
        return {"user_id": 0, "username": "system"}
    return get_current_user(request)


# ---------- Ollama Helpers ----------
def ollama_embed(texts: List[str]) -> List[List[float]]:
    out = []
    for t in texts:
        resp = requests.post(
            f"{OLLAMA_BASE_URL}/api/embeddings",
            json={"model": EMBED_MODEL, "prompt": t},
            timeout=120,
        )
        if resp.status_code != 200:
            raise HTTPException(status_code=500, detail=f"Ollama embeddings error: {resp.text}")
        out.append(resp.json()["embedding"])
    return out


def _extract_chat_content(message_content: Any) -> str:
    """Normalize OpenAI-compatible message content into a plain string."""
    if isinstance(message_content, str):
        return message_content
    if isinstance(message_content, list):
        parts = []
        for item in message_content:
            if isinstance(item, dict):
                if item.get("type") == "text":
                    parts.append(str(item.get("text", "")))
                elif "text" in item:
                    parts.append(str(item["text"]))
            elif item is not None:
                parts.append(str(item))
        return "".join(parts)
    if message_content is None:
        return ""
    return str(message_content)


def ollama_chat(messages: List[Dict[str, str]], extra_body: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Call chat API and return content + performance metrics."""
    if CHAT_API_URL:
        # vLLM / OpenAI-compatible endpoint
        headers = {"Content-Type": "application/json"}
        if CHAT_API_KEY:
            headers["Authorization"] = f"Bearer {CHAT_API_KEY}"
        payload = {"model": CHAT_MODEL, "messages": messages, "stream": False}
        if extra_body:
            payload.update(extra_body)
        t0 = time.time()
        resp = requests.post(
            f"{CHAT_API_URL}/chat/completions",
            headers=headers,
            json=payload,
            timeout=600,
        )
        elapsed = time.time() - t0
        if resp.status_code != 200:
            raise HTTPException(status_code=500, detail=f"Chat API error: {resp.text}")
        data = resp.json()
        try:
            message = data["choices"][0]["message"]
        except (KeyError, IndexError, TypeError) as exc:
            raise HTTPException(status_code=500, detail=f"Unexpected chat API response shape: {exc}")
        content = _extract_chat_content(message.get("content"))
        usage = data.get("usage", {})
        eval_count = usage.get("completion_tokens", 0)
        prompt_tokens = usage.get("prompt_tokens", 0)
        tokens_per_sec = round(eval_count / elapsed, 1) if elapsed > 0 and eval_count else 0
        prompt_tps = round(prompt_tokens / elapsed, 1) if elapsed > 0 and prompt_tokens else 0
        return {
            "content": content,
            "performance": {
                "tokens": eval_count,
                "prompt_tokens": prompt_tokens,
                "tokens_per_sec": tokens_per_sec,
                "prompt_tokens_per_sec": prompt_tps,
                "eval_duration": round(elapsed, 2),
                "total_duration": round(elapsed, 2),
            },
        }

    # Ollama fallback
    resp = requests.post(
        f"{OLLAMA_BASE_URL}/api/chat",
        json={"model": CHAT_MODEL, "messages": messages, "stream": False},
        timeout=600,
    )
    if resp.status_code != 200:
        raise HTTPException(status_code=500, detail=f"Ollama chat error: {resp.text}")
    data = resp.json()
    content = data["message"]["content"]

    # Ollama returns durations in nanoseconds
    eval_count = data.get("eval_count", 0)
    eval_duration_ns = data.get("eval_duration", 0)
    prompt_eval_count = data.get("prompt_eval_count", 0)
    prompt_eval_duration_ns = data.get("prompt_eval_duration", 0)
    total_duration_ns = data.get("total_duration", 0)

    eval_duration_s = eval_duration_ns / 1e9 if eval_duration_ns else 0
    prompt_eval_duration_s = prompt_eval_duration_ns / 1e9 if prompt_eval_duration_ns else 0
    total_duration_s = total_duration_ns / 1e9 if total_duration_ns else 0
    tokens_per_sec = round(eval_count / eval_duration_s, 1) if eval_duration_s > 0 else 0
    prompt_tokens_per_sec = round(prompt_eval_count / prompt_eval_duration_s, 1) if prompt_eval_duration_s > 0 else 0

    return {
        "content": content,
        "performance": {
            "tokens": eval_count,
            "prompt_tokens": prompt_eval_count,
            "tokens_per_sec": tokens_per_sec,
            "prompt_tokens_per_sec": prompt_tokens_per_sec,
            "eval_duration": round(eval_duration_s, 2),
            "total_duration": round(total_duration_s, 2),
        },
    }


# ---------- Text Chunking ----------
_ws = re.compile(r"\s+")


def normalize(text: str) -> str:
    return _ws.sub(" ", text).strip()


def estimate_tokens(text: Optional[str]) -> int:
    if not text:
        return 0
    return max(1, len(text) // APPROX_CHARS_PER_TOKEN)


def trim_text_to_token_budget(text: str, token_budget: int) -> str:
    if token_budget <= 0:
        return ""
    max_chars = token_budget * APPROX_CHARS_PER_TOKEN
    if len(text) <= max_chars:
        return text
    return text[:max(0, max_chars - 3)].rstrip() + "..."


def build_context_from_sources(
    sources: List[Dict[str, Any]],
    token_budget: int,
) -> Tuple[List[Dict[str, Any]], str]:
    selected = []
    blocks = []
    used_tokens = 0
    for source in sources:
        block = f"[{source['source']} #{source['chunk_index']}] {source['text']}"
        block_tokens = estimate_tokens(block)
        if selected and used_tokens + block_tokens > token_budget:
            continue
        if not selected and block_tokens > token_budget:
            block = trim_text_to_token_budget(block, token_budget)
            block_tokens = estimate_tokens(block)
        selected.append(source)
        blocks.append(block)
        used_tokens += block_tokens
        if used_tokens >= token_budget:
            break
    return selected, "\n\n".join(blocks).strip()


def build_web_context(
    results: List[Dict[str, str]],
    token_budget: int,
) -> Tuple[List[Dict[str, Any]], str]:
    selected_sources = []
    blocks = []
    used_tokens = 0
    for result in results:
        block = f"[WEB: {result['title']}]({result['url']})\n{result['content']}"
        block_tokens = estimate_tokens(block)
        if selected_sources and used_tokens + block_tokens > token_budget:
            continue
        if not selected_sources and block_tokens > token_budget:
            block = trim_text_to_token_budget(block, token_budget)
            block_tokens = estimate_tokens(block)
        selected_sources.append({
            "source": result["title"],
            "url": result["url"],
            "score": None,
            "chunk_index": None,
        })
        blocks.append(block)
        used_tokens += block_tokens
        if used_tokens >= token_budget:
            break
    return selected_sources, "\n\n".join(blocks)


def trim_history_messages(history: List[Dict[str, str]], token_budget: int) -> List[Dict[str, str]]:
    kept = []
    used_tokens = 0
    for message in reversed(history):
        content = message.get("content", "")
        message_tokens = estimate_tokens(content)
        if kept and used_tokens + message_tokens > token_budget:
            continue
        if not kept and message_tokens > token_budget:
            message = {**message, "content": trim_text_to_token_budget(content, token_budget)}
            message_tokens = estimate_tokens(message["content"])
        kept.append(message)
        used_tokens += message_tokens
        if used_tokens >= token_budget:
            break
    kept.reverse()
    return kept


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    text = text.strip()
    if not text:
        return []
    chunks = []
    i = 0
    n = len(text)
    while i < n:
        j = min(i + size, n)
        chunks.append(text[i:j])
        if j == n:
            break
        i = max(0, j - overlap)
    return chunks


def list_md_files(root: str) -> List[str]:
    files = []
    if not os.path.isdir(root):
        return files
    for name in os.listdir(root):
        path = os.path.join(root, name)
        if os.path.isfile(path) and name.lower().endswith((".md", ".txt")):
            files.append(path)
    return sorted(files)


def list_ingestible_files(*dirs) -> List[str]:
    """List files in given directories that match TEXT_EXTENSIONS or IMAGE_EXTENSIONS."""
    files = []
    for d in dirs:
        d = Path(d)
        if not d.is_dir():
            continue
        for f in sorted(d.iterdir()):
            if f.is_file() and f.suffix.lower() in TEXT_EXTENSIONS + IMAGE_EXTENSIONS:
                files.append(str(f))
    return files


# ---------- Text Extraction ----------
def extract_text_docx(path: str) -> str:
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_xlsx(path: str) -> str:
    wb = load_workbook(path, read_only=True, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"--- Sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                lines.append("\t".join(cells))
    wb.close()
    return "\n".join(lines)


def extract_text_csv(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        lines = []
        for row in reader:
            if any(cell.strip() for cell in row):
                lines.append("\t".join(row))
    return "\n".join(lines)


def extract_text_pdf(path: str) -> str:
    """Extract text from a PDF file page by page."""
    from PyPDF2 import PdfReader
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages)


def ollama_vision_describe(image_path: str) -> str:
    """Use a vision model (e.g. llava) to describe an image."""
    with open(image_path, "rb") as f:
        img_b64 = base64.b64encode(f.read()).decode("utf-8")
    resp = requests.post(
        f"{OLLAMA_BASE_URL}/api/generate",
        json={
            "model": VISION_MODEL,
            "prompt": "Describe this image in detail. Include all visible text, labels, data, and visual elements.",
            "images": [img_b64],
            "stream": False,
        },
        timeout=300,
    )
    if resp.status_code != 200:
        raise HTTPException(status_code=500, detail=f"Vision model error: {resp.text}")
    return resp.json().get("response", "")


def extract_text_image(path: str) -> str:
    """Extract text from an image using OCR, falling back to vision model."""
    try:
        img = Image.open(path)
        ocr_text = pytesseract.image_to_string(img).strip()
        if len(ocr_text) >= OCR_MIN_CHARS:
            return ocr_text
    except Exception:
        pass
    # Fallback to vision model
    try:
        return ollama_vision_describe(path)
    except Exception:
        return ""


def extract_text_for_file(path: str) -> str:
    """Extract text from a file based on its extension."""
    ext = os.path.splitext(path.lower())[1]
    if ext in (".md", ".txt"):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext == ".docx":
        return extract_text_docx(path)
    elif ext == ".xlsx":
        return extract_text_xlsx(path)
    elif ext == ".csv":
        return extract_text_csv(path)
    elif ext == ".pdf":
        return extract_text_pdf(path)
    elif ext in IMAGE_EXTENSIONS:
        return extract_text_image(path)
    return ""


# ---------- Qdrant Vector Store ----------
def qdrant_client() -> QdrantClient:
    return QdrantClient(url=QDRANT_URL)


def ensure_collection(client: QdrantClient, vector_size: int) -> None:
    existing = [c.name for c in client.get_collections().collections]
    if COLLECTION in existing:
        return
    client.create_collection(
        collection_name=COLLECTION,
        vectors_config=qm.VectorParams(size=vector_size, distance=qm.Distance.COSINE),
    )


def upsert_chunks(
    client: QdrantClient,
    chunks: List[Tuple[str, Dict[str, Any]]],
    vectors: List[List[float]],
) -> int:
    points = []
    for idx, (chunk, meta) in enumerate(chunks):
        points.append(
            qm.PointStruct(
                id=str(uuid.uuid4()),
                vector=vectors[idx],
                payload={"text": chunk, **meta},
            )
        )
    client.upsert(collection_name=COLLECTION, points=points)
    return len(points)


def search(
    client: QdrantClient, query_vec: List[float], top_k: int = TOP_K, project_id: Optional[str] = None
) -> List[Dict[str, Any]]:
    query_filter = None
    if project_id:
        query_filter = qm.Filter(
            must=[qm.FieldCondition(key="project_id", match=qm.MatchValue(value=project_id))]
        )
    response = client.query_points(
        collection_name=COLLECTION,
        query=query_vec,
        limit=top_k,
        query_filter=query_filter,
    )
    results = []
    for h in response.points:
        payload = h.payload or {}
        results.append({
            "score": h.score,
            "text": payload.get("text", ""),
            "source": payload.get("source", ""),
            "chunk_index": payload.get("chunk_index", -1),
            "project_id": payload.get("project_id"),
            "file_id": payload.get("file_id"),
        })
    return results


def delete_points_for_file(client: QdrantClient, file_id: str) -> None:
    client.delete(
        collection_name=COLLECTION,
        points_selector=qm.FilterSelector(
            filter=qm.Filter(
                must=[qm.FieldCondition(key="file_id", match=qm.MatchValue(value=file_id))]
            )
        ),
    )


def build_retrieval_context(
    query: str, top_k: int = TOP_K, project_id: Optional[str] = None
) -> Tuple[List[Dict[str, Any]], str]:
    """Retrieve relevant chunks and format them for prompting."""
    cleaned_query = normalize(query)
    if not cleaned_query:
        return [], ""

    qvec = ollama_embed([cleaned_query])[0]
    client = qdrant_client()
    sources = search(client, qvec, top_k=top_k, project_id=project_id)

    context_blocks = []
    for source in sources:
        if source["text"]:
            context_blocks.append(
                f"[{source['source']} #{source['chunk_index']}] {source['text']}"
            )
    return sources, "\n\n".join(context_blocks).strip()


# ---------- Web Search ----------
_INJECTION_PATTERNS = [
    re.compile(r"(?i)\b(ignore|disregard|override)\b.{0,80}\b(instruction|previous|system|developer|prompt)\b"),
    re.compile(r"(?i)\b(system prompt|developer message|you are chatgpt|assistant:|tool:|function call)\b"),
    re.compile(r"(?i)\b(do not summarize|instead of summarizing|browse to|call the tool|reveal prompt)\b"),
]


def sanitize_web_text(text: str) -> str:
    """Treat fetched web content as untrusted and strip common prompt-injection patterns."""
    cleaned = re.sub(r"<!--.*?-->", " ", text or "", flags=re.DOTALL)
    cleaned = re.sub(r"(?is)<(script|style|noscript|iframe).*?>.*?</\1>", " ", cleaned)
    lines = []
    for raw_line in cleaned.splitlines():
        line = normalize(raw_line)
        if not line:
            continue
        if any(pattern.search(line) for pattern in _INJECTION_PATTERNS):
            continue
        if line.lower().startswith(("javascript:", "data:", "vbscript:")):
            continue
        lines.append(line)
    return normalize("\n".join(lines))[:WEB_SEARCH_PAGE_CHAR_LIMIT]


def clamp_web_results(count: Optional[int]) -> int:
    if count is None:
        return WEB_SEARCH_RESULTS_DEFAULT
    return max(1, min(int(count), WEB_SEARCH_RESULTS_MAX))


def brave_search(query: str, max_results: int) -> List[Dict[str, str]]:
    """Search Brave and return a normalized result list."""
    if not BRAVE_SEARCH_API_KEY:
        raise HTTPException(status_code=503, detail="Web search is not configured")

    params = {
        "q": query,
        "count": clamp_web_results(max_results),
        "extra_snippets": "true",
        "country": WEB_SEARCH_COUNTRY,
        "ui_lang": WEB_SEARCH_UI_LANG,
        "safesearch": BRAVE_SAFESEARCH,
        "spellcheck": "true",
    }
    headers = {
        "Accept": "application/json",
        "X-Subscription-Token": BRAVE_SEARCH_API_KEY,
    }
    try:
        resp = requests.get(BRAVE_SEARCH_BASE_URL, params=params, headers=headers, timeout=30)
        resp.raise_for_status()
    except requests.RequestException as exc:
        raise HTTPException(status_code=502, detail=f"Brave search failed: {exc}")

    data = resp.json()
    results = []
    for item in data.get("web", {}).get("results", []):
        snippets = []
        if item.get("description"):
            snippets.append(item["description"])
        snippets.extend(item.get("extra_snippets") or [])
        snippet_text = normalize(" ".join(snippets))[:WEB_SNIPPET_CHAR_LIMIT]
        results.append({
            "title": item.get("title", "") or urlparse(item.get("url", "")).netloc,
            "url": item.get("url", ""),
            "content": snippet_text,
        })
    return results


def tavily_search(query: str, max_results: int) -> List[Dict[str, str]]:
    """Search Tavily and return a normalized result list."""
    if not TAVILY_API_KEY:
        raise HTTPException(status_code=503, detail="Web search is not configured")

    payload = {
        "api_key": TAVILY_API_KEY,
        "query": query,
        "topic": TAVILY_SEARCH_TOPIC,
        "search_depth": "advanced",
        "max_results": clamp_web_results(max_results),
        "include_answer": False,
        "include_raw_content": False,
        "include_images": False,
    }
    try:
        resp = requests.post(TAVILY_SEARCH_BASE_URL, json=payload, timeout=30)
        resp.raise_for_status()
    except requests.RequestException as exc:
        raise HTTPException(status_code=502, detail=f"Tavily search failed: {exc}")

    data = resp.json()
    results = []
    for item in data.get("results", []):
        url = item.get("url", "")
        snippet_text = normalize(item.get("content", ""))[:WEB_SNIPPET_CHAR_LIMIT]
        results.append({
            "title": item.get("title", "") or urlparse(url).netloc,
            "url": url,
            "content": snippet_text,
        })
    return results


async def _crawl_and_sanitize_results(results: List[Dict[str, str]]) -> List[Dict[str, str]]:
    markdown_generator = DefaultMarkdownGenerator(
        content_filter=PruningContentFilter(threshold=0.5),
        options={"ignore_links": True},
    )
    config = CrawlerRunConfig(
        cache_mode=CacheMode.BYPASS,
        markdown_generator=markdown_generator,
    )

    sanitized = []
    async with AsyncWebCrawler() as crawler:
        for item in results:
            url = item.get("url", "")
            if not url:
                continue
            content = item.get("content", "")
            try:
                crawl_result = await crawler.arun(url=url, config=config)
                if crawl_result.success and crawl_result.markdown:
                    markdown_obj = crawl_result.markdown
                    if hasattr(markdown_obj, "fit_markdown"):
                        content = markdown_obj.fit_markdown or markdown_obj.raw_markdown or content
                    else:
                        content = str(markdown_obj)
            except Exception:
                pass

            sanitized_text = sanitize_web_text(content)
            if not sanitized_text:
                continue
            sanitized.append({
                "title": item.get("title", "") or urlparse(url).netloc,
                "url": url,
                "content": sanitized_text,
            })
    return sanitized


def web_search(query: str, max_results: int = WEB_SEARCH_RESULTS_DEFAULT) -> List[Dict[str, str]]:
    """Search with Tavily or Brave, then crawl and sanitize the resulting pages with Crawl4AI."""
    if TAVILY_API_KEY:
        search_results = tavily_search(query, max_results=max_results)
    elif BRAVE_SEARCH_API_KEY:
        search_results = brave_search(query, max_results=max_results)
    else:
        raise HTTPException(status_code=503, detail="Web search is not configured")

    if not search_results:
        return []
    return asyncio.run(_crawl_and_sanitize_results(search_results))


# ---------- Pydantic Models ----------
class AuthRequest(BaseModel):
    username: str
    password: str


class IngestResponse(BaseModel):
    files: int
    chunks: int
    points_upserted: int


class ProjectCreateRequest(BaseModel):
    name: str


class ProjectRenameRequest(BaseModel):
    name: str


class ProjectAssignRequest(BaseModel):
    project_id: Optional[str] = None


class ChatRequest(BaseModel):
    convo_id: str
    question: str
    project_id: Optional[str] = None
    history_limit: int = 12
    top_k: int = TOP_K
    mode: str = "research"
    web_search: bool = False
    web_results: int = WEB_SEARCH_RESULTS_DEFAULT


class ChatResponse(BaseModel):
    convo_id: str
    answer: str
    sources: List[Dict[str, Any]]
    performance: Dict[str, Any]


class FetchURLRequest(BaseModel):
    url: str
    project_id: str
    filename: Optional[str] = None


def cleanup_exports() -> None:
    """Delete stale export artifacts from the exports directory."""
    now = time.time()
    for path in EXPORTS_DIR.glob("*"):
        try:
            if not path.is_file():
                continue
            age = now - path.stat().st_mtime
            if age > EXPORT_MAX_AGE:
                path.unlink()
        except FileNotFoundError:
            continue
        except OSError as exc:
            print(f"[cleanup_exports] Failed to remove {path}: {exc}")


def sanitize_filename(filename: str) -> str:
    """Reject path-like filenames and keep only a plain file name."""
    cleaned = (filename or "").strip().replace("\\", "/")
    name = Path(cleaned).name.strip()
    if not name or name in {".", ".."}:
        raise HTTPException(status_code=400, detail="Invalid filename")
    if cleaned != name:
        raise HTTPException(status_code=400, detail="Filename must not include path components")
    return name


def unique_project_path(dest_dir: Path, filename: str) -> Tuple[str, Path]:
    stem = Path(filename).stem
    suffix = Path(filename).suffix
    stored_name = filename
    dest_path = dest_dir / stored_name
    if not dest_path.exists():
        return stored_name, dest_path

    counter = 2
    while True:
        stored_name = f"{stem}_{counter}{suffix}"
        dest_path = dest_dir / stored_name
        if not dest_path.exists():
            return stored_name, dest_path
        counter += 1


def safe_export_stem(value: str, fallback: str) -> str:
    stem = re.sub(r"[^\w\s-]", "", value)[:40].strip()
    return stem or fallback


def normalize_chat_mode(mode: str) -> str:
    value = (mode or "research").lower().strip()
    aliases = {
        "hybrid": "research",
        "research": "research",
        "grounded": "grounded",
        "general": "general",
    }
    return aliases.get(value, "research")


# ---------- Health ----------
@app.get("/health")
def health():
    provider = "tavily+crawl4ai" if TAVILY_API_KEY else "brave+crawl4ai"
    return {
        "ok": True,
        "collection": COLLECTION,
        "data_dir": DATA_DIR,
        "web_search_provider": provider,
        "web_search_configured": bool(TAVILY_API_KEY or BRAVE_SEARCH_API_KEY),
    }


# ---------- Auth Endpoints ----------
@app.post("/auth/register")
def auth_register(req: AuthRequest, response: Response):
    username = req.username.strip()
    password = req.password

    if len(username) < 4:
        raise HTTPException(400, "Username must be at least 4 characters")
    if not re.match(r"^[a-zA-Z0-9._-]+$", username):
        raise HTTPException(400, "Username: letters, numbers, dots, underscores, hyphens only")
    if len(password) < 8:
        raise HTTPException(400, "Password must be at least 8 characters")

    hashed = bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
    conn = db()
    try:
        with conn:
            conn.execute(
                "INSERT INTO users (username, password_hash, created_at) VALUES (?, ?, ?)",
                (username, hashed, int(time.time())),
            )
        cur = conn.execute("SELECT id FROM users WHERE username=?", (username,))
        user_id = cur.fetchone()[0]
    except sqlite3.IntegrityError:
        conn.close()
        raise HTTPException(409, "Username already taken")
    conn.close()

    cookie = create_session_cookie(user_id, username)
    response.set_cookie(
        "merlin_session", cookie,
        httponly=True, max_age=SESSION_MAX_AGE, samesite="lax",
        path="/", secure=COOKIE_SECURE,
    )
    return {"ok": True, "user_id": user_id, "username": username}


@app.post("/auth/login")
def auth_login(req: AuthRequest, response: Response):
    username = req.username.strip()
    conn = db()
    cur = conn.execute("SELECT id, password_hash FROM users WHERE username=?", (username,))
    row = cur.fetchone()
    conn.close()

    if not row or not bcrypt.checkpw(req.password.encode("utf-8"), row[1].encode("utf-8")):
        raise HTTPException(401, "Invalid username or password")

    user_id = row[0]
    cookie = create_session_cookie(user_id, username)
    response.set_cookie(
        "merlin_session", cookie,
        httponly=True, max_age=SESSION_MAX_AGE, samesite="lax",
        path="/", secure=COOKIE_SECURE,
    )
    return {"ok": True, "user_id": user_id, "username": username}


@app.post("/auth/logout")
def auth_logout(response: Response):
    response.delete_cookie("merlin_session", path="/")
    return {"ok": True}


@app.get("/auth/me")
def auth_me(user: Dict = Depends(get_current_user)):
    return {"user_id": user["user_id"], "username": user["username"]}


@app.get("/projects")
def list_projects(user: Dict = Depends(get_current_user)):
    conn = db()
    cur = conn.execute(
        "SELECT p.id, p.name, p.created_at, p.updated_at, COUNT(f.id) "
        "FROM projects p "
        "LEFT JOIN project_files f ON f.project_id = p.id "
        "WHERE p.user_id=? "
        "GROUP BY p.id, p.name, p.created_at, p.updated_at "
        "ORDER BY p.updated_at DESC, p.created_at DESC",
        (user["user_id"],),
    )
    rows = cur.fetchall()
    conn.close()
    return [
        {
            "id": r[0],
            "name": r[1],
            "created_at": r[2],
            "updated_at": r[3],
            "file_count": r[4],
        }
        for r in rows
    ]


@app.post("/projects")
def create_project(req: ProjectCreateRequest, user: Dict = Depends(get_current_user)):
    name = req.name.strip()
    if not name:
        raise HTTPException(400, "Project name required")

    project_id = str(uuid.uuid4())
    now = int(time.time())
    conn = db()
    with conn:
        conn.execute(
            "INSERT INTO projects (id, user_id, name, created_at, updated_at) VALUES (?, ?, ?, ?, ?)",
            (project_id, user["user_id"], name[:100], now, now),
        )
    conn.close()
    project_text_dir(user["user_id"], project_id)
    project_incoming_dir(user["user_id"], project_id)
    return {"id": project_id, "name": name[:100], "created_at": now, "updated_at": now, "file_count": 0}


@app.patch("/projects/{project_id}")
def rename_project(project_id: str, req: ProjectRenameRequest, user: Dict = Depends(get_current_user)):
    name = req.name.strip()
    if not name:
        raise HTTPException(400, "Project name required")

    conn = db()
    require_project(conn, project_id, user["user_id"])
    with conn:
        conn.execute(
            "UPDATE projects SET name=?, updated_at=? WHERE id=? AND user_id=?",
            (name[:100], int(time.time()), project_id, user["user_id"]),
        )
    conn.close()
    return {"ok": True}


@app.delete("/projects/{project_id}")
def delete_project(project_id: str, user: Dict = Depends(get_current_user)):
    conn = db()
    require_project(conn, project_id, user["user_id"])
    cur = conn.execute(
        "SELECT id, rel_path FROM project_files WHERE project_id=? AND user_id=?",
        (project_id, user["user_id"]),
    )
    files = cur.fetchall()
    client = qdrant_client()
    for file_id, _ in files:
        delete_points_for_file(client, file_id)

    with conn:
        conn.execute(
            "UPDATE conversations SET project_id=NULL WHERE project_id=? AND user_id=?",
            (project_id, user["user_id"]),
        )
        conn.execute("DELETE FROM project_files WHERE project_id=? AND user_id=?", (project_id, user["user_id"]))
        conn.execute("DELETE FROM projects WHERE id=? AND user_id=?", (project_id, user["user_id"]))
    conn.close()

    shutil.rmtree(user_project_root(user["user_id"], project_id), ignore_errors=True)
    return {"ok": True}


@app.get("/projects/{project_id}/files")
def list_project_files(project_id: str, user: Dict = Depends(get_current_user)):
    conn = db()
    project = require_project(conn, project_id, user["user_id"])
    cur = conn.execute(
        "SELECT id, filename, stored_name, rel_path, media_kind, size, created_at, updated_at "
        "FROM project_files WHERE project_id=? AND user_id=? ORDER BY updated_at DESC, created_at DESC",
        (project_id, user["user_id"]),
    )
    rows = cur.fetchall()
    conn.close()
    return {
        "project": {
            "id": project[0],
            "name": project[2],
            "created_at": project[3],
            "updated_at": project[4],
        },
        "files": [
            {
                "id": r[0],
                "name": r[1],
                "stored_name": r[2],
                "path": r[3],
                "media_kind": r[4],
                "size": r[5],
                "created_at": r[6],
                "updated_at": r[7],
            }
            for r in rows
        ],
    }


@app.delete("/projects/{project_id}/files/{file_id}")
def delete_project_file(project_id: str, file_id: str, user: Dict = Depends(get_current_user)):
    conn = db()
    require_project(conn, project_id, user["user_id"])
    cur = conn.execute(
        "SELECT rel_path FROM project_files WHERE id=? AND project_id=? AND user_id=?",
        (file_id, project_id, user["user_id"]),
    )
    row = cur.fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "File not found")

    delete_points_for_file(qdrant_client(), file_id)
    with conn:
        conn.execute(
            "DELETE FROM project_files WHERE id=? AND project_id=? AND user_id=?",
            (file_id, project_id, user["user_id"]),
        )
        conn.execute(
            "UPDATE projects SET updated_at=? WHERE id=? AND user_id=?",
            (int(time.time()), project_id, user["user_id"]),
        )
    conn.close()

    abs_path = RAG_ROOT / row[0]
    try:
        abs_path.unlink()
    except FileNotFoundError:
        pass
    return {"ok": True}


# ---------- Conversation Endpoints ----------
@app.get("/conversations")
def list_conversations(user: Dict = Depends(get_current_user)):
    conn = db()
    cur = conn.execute(
        "SELECT c.id, c.title, c.created_at, c.updated_at, c.project_id, p.name "
        "FROM conversations c "
        "LEFT JOIN projects p ON p.id = c.project_id "
        "WHERE c.user_id=? ORDER BY c.updated_at DESC",
        (user["user_id"],),
    )
    rows = cur.fetchall()
    conn.close()
    return [
        {
            "id": r[0],
            "title": r[1],
            "created_at": r[2],
            "updated_at": r[3],
            "project_id": r[4],
            "project_name": r[5],
        }
        for r in rows
    ]


@app.post("/conversations")
def create_conversation(req: Optional[ProjectAssignRequest] = None, user: Dict = Depends(get_current_user)):
    conv_id = str(uuid.uuid4())
    now = int(time.time())
    project_id = req.project_id if req else None
    conn = db()
    if project_id:
        require_project(conn, project_id, user["user_id"])
    with conn:
        conn.execute(
            "INSERT INTO conversations (id, user_id, project_id, title, created_at, updated_at) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (conv_id, user["user_id"], project_id, "New Conversation", now, now),
        )
    conn.close()
    return {
        "id": conv_id,
        "title": "New Conversation",
        "created_at": now,
        "updated_at": now,
        "project_id": project_id,
        "project_name": None,
    }


class RenameRequest(BaseModel):
    title: str


@app.patch("/conversations/{conv_id}")
def rename_conversation(conv_id: str, req: RenameRequest, user: Dict = Depends(get_current_user)):
    title = req.title.strip()
    if not title:
        raise HTTPException(400, "Title required")

    conn = db()
    cur = conn.execute("SELECT user_id FROM conversations WHERE id=?", (conv_id,))
    row = cur.fetchone()
    if not row or row[0] != user["user_id"]:
        conn.close()
        raise HTTPException(404, "Conversation not found")

    with conn:
        conn.execute(
            "UPDATE conversations SET title=?, updated_at=? WHERE id=?",
            (title[:100], int(time.time()), conv_id),
        )
    conn.close()
    return {"ok": True}


@app.patch("/conversations/{conv_id}/project")
def assign_conversation_project(conv_id: str, req: ProjectAssignRequest, user: Dict = Depends(get_current_user)):
    conn = db()
    cur = conn.execute("SELECT user_id FROM conversations WHERE id=?", (conv_id,))
    row = cur.fetchone()
    if not row or row[0] != user["user_id"]:
        conn.close()
        raise HTTPException(404, "Conversation not found")

    if req.project_id:
        require_project(conn, req.project_id, user["user_id"])

    with conn:
        conn.execute(
            "UPDATE conversations SET project_id=?, updated_at=? WHERE id=?",
            (req.project_id, int(time.time()), conv_id),
        )
    conn.close()
    return {"ok": True}


@app.delete("/conversations/{conv_id}")
def delete_conversation(conv_id: str, user: Dict = Depends(get_current_user)):
    conn = db()
    cur = conn.execute("SELECT user_id FROM conversations WHERE id=?", (conv_id,))
    row = cur.fetchone()
    if not row or row[0] != user["user_id"]:
        conn.close()
        raise HTTPException(404, "Conversation not found")

    with conn:
        conn.execute("DELETE FROM messages WHERE convo_id=?", (conv_id,))
        conn.execute("DELETE FROM conversations WHERE id=?", (conv_id,))
    conn.close()
    return {"ok": True}


@app.get("/conversations/{conv_id}/messages")
def get_conversation_messages(conv_id: str, user: Dict = Depends(get_current_user)):
    conn = db()
    cur = conn.execute("SELECT user_id, project_id FROM conversations WHERE id=?", (conv_id,))
    row = cur.fetchone()
    if not row or row[0] != user["user_id"]:
        conn.close()
        raise HTTPException(404, "Conversation not found")

    cur = conn.execute(
        "SELECT role, content, ts FROM messages WHERE convo_id=? ORDER BY ts ASC, id ASC",
        (conv_id,),
    )
    rows = cur.fetchall()
    conn.close()
    return {
        "conversation": {"id": conv_id, "project_id": row[1]},
        "messages": [{"role": r[0], "content": r[1], "ts": r[2]} for r in rows],
    }


# ---------- Chat Endpoint ----------
@app.post("/chat", response_model=ChatResponse)
def chat(req: ChatRequest, request: Request):
    # Authenticate: internal key or session cookie
    user = None
    internal = request.headers.get("X-Internal-Key", "")
    if INTERNAL_KEY and hmac.compare_digest(internal, INTERNAL_KEY):
        user = {"user_id": 0, "username": "system"}
    else:
        cookie = request.cookies.get("merlin_session")
        if cookie:
            user = verify_session_cookie(cookie)
        if not user:
            raise HTTPException(401, "Not authenticated")

    # Ensure conversation exists
    conn = db()
    cur = conn.execute(
        "SELECT id, user_id, title, project_id FROM conversations WHERE id=?", (req.convo_id,)
    )
    conv_row = cur.fetchone()

    if conv_row:
        if user["user_id"] != 0 and conv_row[1] != user["user_id"]:
            conn.close()
            raise HTTPException(403, "Access denied")
        conv_title = conv_row[2]
        conversation_project_id = conv_row[3]
    else:
        now = int(time.time())
        if req.project_id and user["user_id"] != 0:
            require_project(conn, req.project_id, user["user_id"])
        with conn:
            conn.execute(
                "INSERT INTO conversations (id, user_id, project_id, title, created_at, updated_at) "
                "VALUES (?, ?, ?, ?, ?, ?)",
                (req.convo_id, user["user_id"], req.project_id, "New Conversation", now, now),
            )
        conv_title = "New Conversation"
        conversation_project_id = req.project_id
    if req.project_id and req.project_id != conversation_project_id:
        if user["user_id"] != 0:
            require_project(conn, req.project_id, user["user_id"])
        with conn:
            conn.execute(
                "UPDATE conversations SET project_id=?, updated_at=? WHERE id=?",
                (req.project_id, int(time.time()), req.convo_id),
            )
        conversation_project_id = req.project_id
    conn.close()

    mode = normalize_chat_mode(req.mode)
    use_doc_retrieval = mode in ("grounded", "research")
    # Research mode always includes web search; other modes do not.
    use_web_search = mode == "research"

    # Build corpus inventory for system awareness
    conn = db()
    corpus_names = []
    if conversation_project_id:
        cur = conn.execute(
            "SELECT filename FROM project_files WHERE project_id=? ORDER BY updated_at DESC, created_at DESC",
            (conversation_project_id,),
        )
        corpus_names = [row[0] for row in cur.fetchall()]
    conn.close()
    corpus_summary = ", ".join(corpus_names) if corpus_names else "(no files ingested in project)"

    sources = []
    doc_sources = []
    context = ""
    has_enough_context = False
    if use_doc_retrieval:
        qvec = ollama_embed([req.question])[0]
        client = qdrant_client()
        sources = search(client, qvec, top_k=req.top_k, project_id=conversation_project_id)
        top_score = sources[0]["score"] if sources else 0.0
        filtered_sources = [s for s in sources if s["text"] and (s["score"] or 0.0) >= MIN_TOP_SCORE]
        doc_sources, context = build_context_from_sources(
            filtered_sources,
            token_budget=DOC_CONTEXT_TOKEN_BUDGET,
        )
        has_enough_context = (len(context) >= MIN_CONTEXT_CHARS) or (top_score >= MIN_TOP_SCORE)

    # Web search (if enabled by user)
    web_context = ""
    web_results = []
    web_sources = []
    if use_web_search:
        web_results = web_search(req.question, max_results=req.web_results)
        if web_results:
            web_sources, web_context = build_web_context(
                web_results,
                token_budget=WEB_CONTEXT_TOKEN_BUDGET,
            )

    system = (
        "You are a helpful assistant.\n"
        f"CORPUS INVENTORY ({len(corpus_names)} files): {corpus_summary}\n"
        "You must answer in two sections:\n"
        "1) DOCUMENT-BASED ANALYSIS: Use the provided CONTEXT first. "
        "Cite sources like [filename #chunk].\n"
    )
    if web_context:
        system += (
            "2) WEB SEARCH RESULTS: Summarize relevant findings from WEB CONTEXT. "
            "Cite sources like [WEB: title](url).\n"
            "3) GENERAL KNOWLEDGE: If the context and web results do not fully cover "
            "the question, add a third section using your own training knowledge.\n"
            "When WEB CONTEXT is present, prioritize it for current events, dates, schedules, scores, and other time-sensitive facts.\n"
        )
    else:
        system += (
            "2) GENERAL KNOWLEDGE: If the context does not fully cover the question, "
            "add a second section using your own training knowledge.\n"
        )
    system += (
        "   In this section, do NOT invent citations to files. "
        "Clearly label it as general knowledge.\n"
        "Rules:\n"
        "- Never claim you read files that are not in CONTEXT.\n"
        "- If the user asks about what files or documents are available, "
        "refer to the CORPUS INVENTORY above.\n"
        "- If the user asks for something not in the documents, "
        "say so in section 1, then proceed in section 2.\n"
        "- If mode=grounded, ONLY produce section 1 and say what is missing.\n"
    )

    history = trim_history_messages(
        get_recent_messages(req.convo_id, limit=req.history_limit),
        token_budget=HISTORY_TOKEN_BUDGET,
    )
    add_message(req.convo_id, "user", req.question, user["user_id"])
    messages = [{"role": "system", "content": system}]

    for m in history:
        if m["role"] in ("user", "assistant"):
            messages.append(m)

    web_section = ""
    if web_context:
        web_section = f"\n\nWEB CONTEXT:\n{web_context}"

    if mode == "general":
        content = req.question
        if web_context:
            content = f"WEB CONTEXT:\n{web_context}\n\nQUESTION:\n{req.question}"
        messages.append({"role": "user", "content": content})
    elif mode == "grounded":
        ctx = context if context else "(no retrieved context)"
        messages.append({
            "role": "user",
            "content": f"CONTEXT:\n{ctx}{web_section}\n\nQUESTION:\n{req.question}\n\nMODE: grounded (docs only).",
        })
    else:
        ctx = context if context else "(no retrieved context)"
        coverage_note = (
            "Retrieved documents appear relevant."
            if has_enough_context
            else "Retrieved documents appear limited; rely primarily on WEB CONTEXT and clearly separate any general knowledge."
        )
        messages.append({
            "role": "user",
            "content": f"CONTEXT:\n{ctx}{web_section}\n\nQUESTION:\n{req.question}\n\nMODE: research (documents + web). {coverage_note}",
        })

    prompt_token_estimate = sum(estimate_tokens(message.get("content", "")) for message in messages)
    max_prompt_tokens = max(1024, MODEL_MAX_INPUT_TOKENS - MODEL_OUTPUT_TOKEN_RESERVE)
    if prompt_token_estimate > max_prompt_tokens:
        raise HTTPException(
            status_code=400,
            detail=(
                "The requested analysis is too large for the current model context window. "
                "Try fewer files, lower detail, or a smaller history window."
            ),
        )

    result = ollama_chat(messages)
    answer = result["content"]
    add_message(req.convo_id, "assistant", answer, user["user_id"])

    # Auto-title and update timestamp
    conn = db()
    with conn:
        if conv_title == "New Conversation":
            title = req.question[:50].strip()
            if len(req.question) > 50:
                title += "..."
            conn.execute(
                "UPDATE conversations SET title=?, updated_at=? WHERE id=?",
                (title, int(time.time()), req.convo_id),
            )
        else:
            conn.execute(
                "UPDATE conversations SET updated_at=? WHERE id=?",
                (int(time.time()), req.convo_id),
            )
    conn.close()

    return ChatResponse(
        convo_id=req.convo_id,
        answer=answer,
        sources=doc_sources + web_sources,
        performance=result["performance"],
    )


# ---------- File Endpoints ----------
@app.get("/files")
def files(request: Request, user: Dict = Depends(get_optional_user)):
    project_id = (request.query_params.get("project_id") or "").strip()
    if project_id:
        conn = db()
        require_project(conn, project_id, user["user_id"])
        cur = conn.execute(
            "SELECT id, filename, stored_name, rel_path, media_kind, size, created_at, updated_at "
            "FROM project_files WHERE project_id=? AND user_id=? ORDER BY updated_at DESC, created_at DESC",
            (project_id, user["user_id"]),
        )
        rows = cur.fetchall()
        conn.close()
        return {
            "project_id": project_id,
            "text": [
                {
                    "id": r[0],
                    "name": r[1],
                    "stored_name": r[2],
                    "path": r[3],
                    "media_kind": r[4],
                    "size": r[5],
                    "created_at": r[6],
                    "updated_at": r[7],
                }
                for r in rows
            ],
            "incoming": [],
        }

    def list_dir(p: Path):
        if not p.exists():
            return []
        out = []
        for f in sorted(p.glob("*")):
            if f.is_file():
                out.append({"name": f.name, "size": f.stat().st_size, "path": str(f), "media_kind": "legacy"})
        return out

    return {"project_id": None, "text": list_dir(TEXT_DIR), "incoming": list_dir(INCOMING_DIR)}


def _ingest_single_file(filepath: str, file_id: str, project_id: str) -> None:
    """Background task: extract text, chunk, embed, and upsert a single file."""
    try:
        text = extract_text_for_file(filepath)
        if not text or not text.strip():
            return
        text = normalize(text)
        pieces = chunk_text(text)
        if not pieces:
            return
        chunks = [
            (
                p,
                {
                    "source": os.path.basename(filepath),
                    "chunk_index": i,
                    "project_id": project_id,
                    "file_id": file_id,
                },
            )
            for i, p in enumerate(pieces)
        ]
        vecs = ollama_embed([c[0] for c in chunks])
        client = qdrant_client()
        ensure_collection(client, vector_size=len(vecs[0]))
        delete_points_for_file(client, file_id)
        upsert_chunks(client, chunks, vecs)
    except Exception as e:
        print(f"[_ingest_single_file] Error processing {filepath}: {e}")


@app.post("/upload")
async def upload(
    file: UploadFile = File(...),
    project_id: str = Form(...),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    user: Dict = Depends(get_optional_user),
):
    project_id = (project_id or "").strip()
    if not project_id:
        raise HTTPException(status_code=400, detail="project_id is required")

    conn = db()
    require_project(conn, project_id, user["user_id"])
    filename = sanitize_filename(file.filename or "")

    ext = os.path.splitext(filename.lower())[1]
    is_text = ext in TEXT_EXTENSIONS
    is_image = ext in IMAGE_EXTENSIONS
    media_kind = "text" if is_text else ("image" if is_image else "incoming")
    dest_dir = project_text_dir(user["user_id"], project_id) if is_text else project_incoming_dir(user["user_id"], project_id)
    stored_name, dest_path = unique_project_path(dest_dir, filename)

    raw = await file.read()
    dest_path.write_bytes(raw)
    now = int(time.time())
    file_id = str(uuid.uuid4())
    rel_path = str(dest_path.relative_to(RAG_ROOT))
    with conn:
        conn.execute(
            "INSERT INTO project_files (id, project_id, user_id, filename, stored_name, rel_path, media_kind, size, created_at, updated_at) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (file_id, project_id, user["user_id"], filename, stored_name, rel_path, media_kind, len(raw), now, now),
        )
        conn.execute(
            "UPDATE projects SET updated_at=? WHERE id=? AND user_id=?",
            (now, project_id, user["user_id"]),
        )
    conn.close()

    ingesting = is_text or is_image
    if ingesting:
        background_tasks.add_task(_ingest_single_file, str(dest_path), file_id, project_id)

    return {
        "ok": True,
        "file_id": file_id,
        "saved_to": str(dest_path),
        "project_id": project_id,
        "needs_conversion": not (is_text or is_image),
        "ingesting": ingesting,
    }


@app.post("/fetch-url")
def fetch_url(req: FetchURLRequest, background_tasks: BackgroundTasks = BackgroundTasks(), user: Dict = Depends(get_optional_user)):
    """Download a file from a URL, save it, and auto-ingest if possible."""
    url = req.url.strip()
    if not url:
        raise HTTPException(status_code=400, detail="Missing url")
    project_id = (req.project_id or "").strip()
    if not project_id:
        raise HTTPException(status_code=400, detail="project_id is required")

    conn = db()
    require_project(conn, project_id, user["user_id"])

    try:
        resp = requests.get(url, timeout=120, stream=True, headers={"User-Agent": "MerlinRAG/1.0"})
        resp.raise_for_status()
    except requests.RequestException as e:
        raise HTTPException(status_code=502, detail=f"Failed to fetch URL: {e}")

    # Determine filename
    filename = (req.filename or "").strip()
    if not filename:
        cd = resp.headers.get("Content-Disposition", "")
        if "filename=" in cd:
            filename = cd.split("filename=")[-1].strip().strip('"').strip("'")
        if not filename:
            filename = unquote(urlparse(url).path.split("/")[-1]) or "downloaded_file"
    filename = sanitize_filename(filename)

    ext = os.path.splitext(filename.lower())[1]
    is_text = ext in TEXT_EXTENSIONS
    is_image = ext in IMAGE_EXTENSIONS
    media_kind = "text" if is_text else ("image" if is_image else "incoming")
    dest_dir = project_text_dir(user["user_id"], project_id) if is_text else project_incoming_dir(user["user_id"], project_id)
    stored_name, dest_path = unique_project_path(dest_dir, filename)

    dest_path.write_bytes(resp.content)
    now = int(time.time())
    file_id = str(uuid.uuid4())
    rel_path = str(dest_path.relative_to(RAG_ROOT))
    with conn:
        conn.execute(
            "INSERT INTO project_files (id, project_id, user_id, filename, stored_name, rel_path, media_kind, size, created_at, updated_at) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (file_id, project_id, user["user_id"], filename, stored_name, rel_path, media_kind, len(resp.content), now, now),
        )
        conn.execute(
            "UPDATE projects SET updated_at=? WHERE id=? AND user_id=?",
            (now, project_id, user["user_id"]),
        )
    conn.close()

    ingesting = is_text or is_image
    if ingesting:
        background_tasks.add_task(_ingest_single_file, str(dest_path), file_id, project_id)

    return {
        "ok": True,
        "file_id": file_id,
        "saved_to": str(dest_path),
        "project_id": project_id,
        "needs_conversion": not (is_text or is_image),
        "ingesting": ingesting,
    }


@app.post("/ingest", response_model=IngestResponse)
def ingest(request: ProjectAssignRequest, user: Dict = Depends(get_optional_user)):
    if not request.project_id:
        raise HTTPException(status_code=400, detail="project_id is required")

    conn = db()
    require_project(conn, request.project_id, user["user_id"])
    cur = conn.execute(
        "SELECT id, rel_path, filename FROM project_files WHERE project_id=? AND user_id=? AND media_kind IN ('text', 'image')",
        (request.project_id, user["user_id"]),
    )
    file_rows = cur.fetchall()
    conn.close()
    if not file_rows:
        raise HTTPException(status_code=400, detail="No project files found to ingest")

    chunks: List[Tuple[str, Dict[str, Any]]] = []
    client = qdrant_client()
    ingested_files = 0
    for file_id, rel_path, filename in file_rows:
        abs_path = RAG_ROOT / rel_path
        content = extract_text_for_file(str(abs_path))
        content = normalize(content or "")
        if not content.strip():
            continue
        pieces = chunk_text(content)
        for i, p in enumerate(pieces):
            chunks.append((
                p,
                {
                    "source": filename,
                    "chunk_index": i,
                    "project_id": request.project_id,
                    "file_id": file_id,
                },
            ))
        delete_points_for_file(client, file_id)
        ingested_files += 1

    if not chunks:
        raise HTTPException(status_code=400, detail="No chunks produced (files empty?)")

    first_vec = ollama_embed([chunks[0][0]])[0]
    ensure_collection(client, vector_size=len(first_vec))

    texts = [c[0] for c in chunks]
    vecs = ollama_embed(texts)
    upserted = upsert_chunks(client, chunks, vecs)
    return IngestResponse(files=ingested_files, chunks=len(chunks), points_upserted=upserted)


# ---------- Export Endpoints ----------
@app.get("/export/{conv_id}")
def export_conversation(conv_id: str, format: str = "md", user: Dict = Depends(get_current_user)):
    """Export a conversation as markdown, JSON, TXT, CSV, or DOCX."""
    conn = db()
    cur = conn.execute("SELECT user_id, title FROM conversations WHERE id=?", (conv_id,))
    row = cur.fetchone()
    if not row or row[0] != user["user_id"]:
        conn.close()
        raise HTTPException(404, "Conversation not found")
    title = row[1]

    cur = conn.execute(
        "SELECT role, content, ts FROM messages WHERE convo_id=? ORDER BY ts ASC, id ASC",
        (conv_id,),
    )
    msgs = [{"role": r[0], "content": r[1], "ts": r[2]} for r in cur.fetchall()]
    conn.close()

    export_format = (format or "md").lower().strip()
    if export_format not in {"md", "json", "txt", "csv", "docx"}:
        raise HTTPException(400, "Unsupported export format")

    safe_title = safe_export_stem(title, "conversation")
    cleanup_exports()

    if export_format == "json":
        export_data = {"id": conv_id, "title": title, "messages": msgs}
        filename = f"{safe_title}_{conv_id[:8]}.json"
        dest = EXPORTS_DIR / filename
        dest.write_text(json.dumps(export_data, indent=2), encoding="utf-8")
        return FileResponse(str(dest), filename=filename, media_type="application/json")

    lines = [f"# {title}\n"]
    for m in msgs:
        ts_str = datetime.fromtimestamp(m["ts"]).strftime("%Y-%m-%d %H:%M:%S") if m["ts"] else ""
        role_label = "**You**" if m["role"] == "user" else "**Merlin**"
        lines.append(f"### {role_label} ({ts_str})\n\n{m['content']}\n\n---\n")

    if export_format == "md":
        filename = f"{safe_title}_{conv_id[:8]}.md"
        dest = EXPORTS_DIR / filename
        dest.write_text("\n".join(lines), encoding="utf-8")
        return FileResponse(str(dest), filename=filename, media_type="text/markdown")

    if export_format == "txt":
        txt_lines = []
        for m in msgs:
            ts_str = datetime.fromtimestamp(m["ts"]).strftime("%Y-%m-%d %H:%M:%S") if m["ts"] else ""
            role_label = "You" if m["role"] == "user" else "Merlin"
            txt_lines.append(f"{role_label} ({ts_str})")
            txt_lines.append(m["content"])
            txt_lines.append("")
        filename = f"{safe_title}_{conv_id[:8]}.txt"
        dest = EXPORTS_DIR / filename
        dest.write_text("\n".join(txt_lines), encoding="utf-8")
        return FileResponse(str(dest), filename=filename, media_type="text/plain")

    if export_format == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["role", "timestamp", "content"])
        for m in msgs:
            ts_str = datetime.fromtimestamp(m["ts"]).strftime("%Y-%m-%d %H:%M:%S") if m["ts"] else ""
            writer.writerow([m["role"], ts_str, m["content"]])
        filename = f"{safe_title}_{conv_id[:8]}.csv"
        dest = EXPORTS_DIR / filename
        dest.write_text(output.getvalue(), encoding="utf-8", newline="")
        return FileResponse(str(dest), filename=filename, media_type="text/csv")

    filename = f"{safe_title}_{conv_id[:8]}.docx"
    dest = EXPORTS_DIR / filename
    markdown_to_docx("\n".join(lines)).save(str(dest))
    media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return FileResponse(str(dest), filename=filename, media_type=media)


class FileGenerateRequest(BaseModel):
    convo_id: Optional[str] = None
    prompt: str
    format: str = "txt"  # csv, txt, docx, png
    context_messages: int = 10


def generate_file_content(
    prompt: str,
    fmt: str,
    context: List[Dict[str, str]],
    rag_context: str = "",
) -> Dict:
    """Ask LLM to generate file content as structured JSON."""
    context_sections = []
    if context:
        context_text = "Conversation context:\n"
        for m in context:
            context_text += f"{m['role']}: {m['content']}\n"
        context_sections.append(context_text.strip())

    if rag_context:
        context_sections.append(f"Retrieved document context:\n{rag_context}")

    context_text = ""
    if context_sections:
        context_text = "\n\nSupporting context:\n" + "\n\n".join(context_sections)

    if fmt == "png":
        system_msg = (
            "You generate chart specifications as JSON. "
            "Return ONLY valid JSON with no markdown fences. "
            "Use the supporting context to populate the chart with actual values when available. "
            'Format: {"chart_type": "bar|line|pie|scatter", "title": "...", '
            '"data": {"labels": [...], "values": [...]}, "xlabel": "...", "ylabel": "...", '
            '"filename_hint": "..."}'
        )
    else:
        system_msg = (
            "You generate file content as JSON. "
            "Return ONLY valid JSON with no markdown fences. "
            "Use the supporting context to produce substantive content grounded in the provided documents when available. "
            "Do not return an empty content field. "
            f'Format: {{"content": "the {fmt} content as a string", "filename_hint": "short_name"}}'
        )
        if fmt == "csv":
            system_msg += "\nFor CSV, the content must be valid CSV text with headers and at least one data row."

    messages = [
        {"role": "system", "content": system_msg},
        {"role": "user", "content": prompt + context_text},
    ]
    extra_body = {"response_format": {"type": "json_object"}} if CHAT_API_URL else None
    raw = ollama_chat(messages, extra_body=extra_body)["content"]
    raw = raw.strip()

    # Strip markdown fences
    fence_match = re.search(r'```(?:json)?\s*\n?(.*?)```', raw, re.DOTALL)
    if fence_match:
        raw = fence_match.group(1).strip()

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        brace_match = re.search(r'\{.*\}', raw, re.DOTALL)
        if brace_match:
            parsed = json.loads(brace_match.group(0))
        else:
            raise
    if not isinstance(parsed, dict):
        raise HTTPException(status_code=500, detail="Generated file payload was not a JSON object")
    return parsed


def render_chart_png(spec: Dict) -> Path:
    """Render a chart from a specification dict using matplotlib."""
    cleanup_exports()
    chart_type = spec.get("chart_type", "bar")
    title = spec.get("title", "Chart")
    data = spec.get("data", {})
    labels = data.get("labels", [])
    values = data.get("values", [])
    xlabel = spec.get("xlabel", "")
    ylabel = spec.get("ylabel", "")
    hint = spec.get("filename_hint", "chart")
    safe_hint = re.sub(r'[^\w\s-]', '', hint)[:30].strip() or "chart"

    fig, ax = plt.subplots(figsize=(10, 6))
    if chart_type == "pie":
        ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=90)
        ax.axis('equal')
    elif chart_type == "line":
        ax.plot(labels, values, marker='o', linewidth=2)
        ax.set_xlabel(xlabel); ax.set_ylabel(ylabel)
    elif chart_type == "scatter":
        ax.scatter(labels, values)
        ax.set_xlabel(xlabel); ax.set_ylabel(ylabel)
    else:
        ax.bar(labels, values)
        ax.set_xlabel(xlabel); ax.set_ylabel(ylabel)
    ax.set_title(title)
    plt.tight_layout()

    filepath = EXPORTS_DIR / f"{safe_hint}_{int(time.time())}.png"
    fig.savefig(str(filepath), dpi=150, bbox_inches='tight')
    plt.close(fig)
    return filepath


def markdown_to_docx(text: str) -> DocxDocument:
    """Convert basic markdown text to a DOCX document."""
    doc = DocxDocument()
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ", "+ ")):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', stripped):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', stripped), style='List Number')
        else:
            doc.add_paragraph(stripped)
    return doc


@app.post("/generate-file")
def generate_file(req: FileGenerateRequest, user: Dict = Depends(get_current_user)):
    fmt = req.format.lower().strip()
    if fmt not in ("csv", "txt", "docx", "png"):
        raise HTTPException(400, "Format must be csv, txt, docx, or png")

    context = []
    project_id = None
    if req.convo_id:
        conn = db()
        cur = conn.execute("SELECT user_id, project_id FROM conversations WHERE id=?", (req.convo_id,))
        row = cur.fetchone()
        if not row or row[0] != user["user_id"]:
            conn.close()
            raise HTTPException(404, "Conversation not found")
        project_id = row[1]
        cur = conn.execute(
            "SELECT role, content FROM messages WHERE convo_id=? ORDER BY ts DESC, id DESC LIMIT ?",
            (req.convo_id, req.context_messages),
        )
        rows = cur.fetchall()
        conn.close()
        context = [{"role": r[0], "content": r[1]} for r in reversed(rows)]

    retrieval_query_parts = [req.prompt]
    for message in context:
        if message["content"]:
            retrieval_query_parts.append(message["content"])
    retrieval_query = "\n".join(retrieval_query_parts)
    _, rag_context = build_retrieval_context(
        retrieval_query,
        top_k=min(req.context_messages, TOP_K),
        project_id=project_id,
    )

    spec = None
    for attempt in range(3):
        try:
            spec = generate_file_content(req.prompt, fmt, context, rag_context=rag_context)
            break
        except json.JSONDecodeError as e:
            print(f"[generate-file] JSON parse attempt {attempt+1} failed: {e}")
            if attempt < 2:
                continue
            raise HTTPException(500, f"Failed to generate valid file content from LLM: {e}")

    if spec is None:
        raise HTTPException(500, "Failed to generate file content")

    cleanup_exports()

    try:
        if fmt == "png":
            filepath = render_chart_png(spec)
            return FileResponse(str(filepath), media_type="image/png", filename=filepath.name,
                                headers={"Content-Disposition": f'attachment; filename="{filepath.name}"'})

        content = spec.get("content", "")
        if content is None:
            content = ""
        if not isinstance(content, str):
            content = json.dumps(content, indent=2)
        content = content.strip()
        if fmt != "png" and not content:
            raise HTTPException(
                status_code=500,
                detail="File generation returned empty content. Try a more specific prompt or verify relevant documents are ingested.",
            )

        hint = spec.get("filename_hint", "generated")
        if not isinstance(hint, str):
            hint = str(hint)
        safe_hint = re.sub(r'[^\w\s-]', '', hint)[:30].strip() or "generated"

        if fmt == "csv":
            filepath = EXPORTS_DIR / f"{safe_hint}_{int(time.time())}.csv"
            filepath.write_text(content, encoding="utf-8")
            return FileResponse(str(filepath), media_type="text/csv", filename=filepath.name,
                                headers={"Content-Disposition": f'attachment; filename="{filepath.name}"'})
        if fmt == "docx":
            filepath = EXPORTS_DIR / f"{safe_hint}_{int(time.time())}.docx"
            doc = markdown_to_docx(content)
            doc.save(str(filepath))
            media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            return FileResponse(str(filepath), media_type=media, filename=filepath.name,
                                headers={"Content-Disposition": f'attachment; filename="{filepath.name}"'})

        filepath = EXPORTS_DIR / f"{safe_hint}_{int(time.time())}.txt"
        filepath.write_text(content, encoding="utf-8")
        return FileResponse(str(filepath), media_type="text/plain", filename=filepath.name,
                            headers={"Content-Disposition": f'attachment; filename="{filepath.name}"'})
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"File generation failed: {exc}")


# ---------- UI ----------
@app.get("/", response_class=HTMLResponse)
def ui():
    return """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>MERLIN &mdash; Knowledge Retrieval System</title>
<link href="https://fonts.googleapis.com/css2?family=Crimson+Text:ital,wght@0,400;0,600;0,700;1,400&display=swap" rel="stylesheet"/>
<style>
:root{
  --bg:#0f1a14;
  --panel-dark:#132018;
  --panel:#1b2a20;
  --surface:#223527;
  --hover:#2f4a37;
  --scroll:#3c5a44;
  --gold-bright:#e0c26a;
  --gold-warm:#c2a55f;
  --gold-dim:#9b7f3c;
  --gold-glow:rgba(224,194,106,0.16);
  --crimson:#7a2b2b;
  --crimson-soft:#a24a4a;
  --verdigris:#4f8f6b;
  --text:#e7e3d7;
  --text-sec:#b6b1a5;
  --text-dim:#8a8577;
  --text-bright:#f6f2e6;
  --parchment:#f3ead6;
  --border:rgba(194,165,95,0.14);
  --border-sub:rgba(205,187,155,0.08);
  --border-orn:rgba(224,194,106,0.26);
  --shadow:0 10px 36px rgba(0,0,0,0.62);
  --shadow-sm:0 3px 14px rgba(0,0,0,0.46);
  --shadow-glow:0 0 26px rgba(224,194,106,0.10);
  --danger:#a24a4a;
  --success:#4f8f6b;
  --radius:4px;
  --font-h:'Crimson Text','Palatino Linotype','Book Antiqua',Georgia,serif;
  --font-ui:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;
  --font-mono:'SF Mono','Cascadia Code',Consolas,'Liberation Mono',monospace;
  --tr:0.2s ease;
}
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
html{font-size:15px;-webkit-font-smoothing:antialiased}
body{font-family:var(--font-h);background:var(--bg);color:var(--text);line-height:1.6;min-height:100vh}
a{color:var(--gold-warm);text-decoration:none}
a:hover{color:var(--gold-bright)}
.hidden{display:none!important}

/* ---- Login ---- */
.login-view{
  display:flex;align-items:center;justify-content:center;min-height:100vh;
  background:radial-gradient(ellipse 800px 600px at 50% 35%,rgba(212,168,67,0.04),transparent),
             radial-gradient(ellipse 400px 300px at 50% 50%,rgba(30,27,24,0.8),transparent),
             var(--bg);
}
.login-card{
  width:100%;max-width:420px;padding:52px 44px 44px;
  background:linear-gradient(180deg,rgba(30,27,24,0.95),rgba(20,18,16,0.98));
  border:1px solid var(--border-orn);border-top:2px solid var(--gold-dim);
  box-shadow:var(--shadow),var(--shadow-glow);position:relative;
}
.login-card::before,.login-card::after{
  content:'';position:absolute;width:20px;height:20px;border-color:var(--gold-dim);opacity:0.5;
}
.login-card::before{top:8px;left:8px;border-top:1px solid;border-left:1px solid}
.login-card::after{top:8px;right:8px;border-top:1px solid;border-right:1px solid}
.corner-bl,.corner-br{position:absolute;width:20px;height:20px;border-color:var(--gold-dim);opacity:0.5}
.corner-bl{bottom:8px;left:8px;border-bottom:1px solid;border-left:1px solid}
.corner-br{bottom:8px;right:8px;border-bottom:1px solid;border-right:1px solid}
.login-brand{text-align:center;margin-bottom:36px}
.login-crest{width:56px;height:56px;margin:0 auto 16px;color:var(--gold-warm);opacity:0.8}
.login-wordmark{
  font-family:var(--font-h);font-size:2.6rem;font-weight:700;
  letter-spacing:0.28em;color:var(--parchment);
  text-shadow:0 0 40px rgba(212,168,67,0.12);
}
.login-accent{width:80px;height:1px;background:linear-gradient(90deg,transparent,var(--gold-warm),transparent);margin:14px auto}
.login-tagline{font-family:var(--font-h);font-size:0.85rem;font-style:italic;color:var(--text-sec);letter-spacing:0.06em}

.auth-tabs{display:flex;gap:0;margin-bottom:28px;border-bottom:1px solid var(--border)}
.auth-tab{
  flex:1;padding:10px;background:none;border:none;border-bottom:2px solid transparent;
  font-family:var(--font-h);font-size:0.95rem;color:var(--text-dim);cursor:pointer;
  transition:all var(--tr);
}
.auth-tab:hover{color:var(--text-sec)}
.auth-tab.active{color:var(--gold-warm);border-bottom-color:var(--gold-warm)}

.form-group{margin-bottom:20px}
.form-label{
  display:block;font-family:var(--font-ui);font-size:0.7rem;font-weight:600;
  color:var(--text-dim);text-transform:uppercase;letter-spacing:0.1em;margin-bottom:6px;
}
.form-input{
  width:100%;padding:11px 14px;background:var(--panel-dark);
  border:1px solid var(--border);border-bottom:1px solid var(--border-orn);
  color:var(--text);font-family:var(--font-h);font-size:1rem;
  outline:none;transition:border-color var(--tr),box-shadow var(--tr);
}
.form-input:focus{border-color:var(--gold-dim);box-shadow:0 1px 0 0 var(--gold-dim)}
.form-input::placeholder{color:var(--text-dim);font-style:italic}
.login-error{font-family:var(--font-ui);font-size:0.8rem;color:var(--crimson-soft);min-height:1.3em;margin-bottom:8px}

/* ---- Buttons ---- */
.btn{
  display:inline-flex;align-items:center;justify-content:center;gap:6px;
  padding:10px 22px;font-family:var(--font-ui);font-size:0.8rem;font-weight:600;
  letter-spacing:0.04em;border:1px solid var(--border);background:var(--surface);
  color:var(--text);cursor:pointer;transition:all var(--tr);white-space:nowrap;
  text-transform:uppercase;border-radius:var(--radius);
}
.btn:hover{background:var(--hover);border-color:var(--border-orn)}
.btn:disabled{opacity:0.4;cursor:not-allowed}
.btn--primary{
  background:linear-gradient(180deg,var(--gold-warm),var(--gold-dim));
  color:var(--bg);border:1px solid var(--gold-warm);
  position:relative;overflow:hidden;
}
.btn--primary:hover{
  background:linear-gradient(180deg,var(--gold-bright),var(--gold-warm));
  box-shadow:0 0 12px rgba(224,194,106,0.35),0 0 30px rgba(224,194,106,0.15);
  transform:translateY(-1px);
}
.btn--primary::after{
  content:'';position:absolute;top:0;left:-75%;width:50%;height:100%;
  background:linear-gradient(120deg,transparent,rgba(255,255,255,0.4),transparent);
  transform:skewX(-20deg);
}
.btn--primary:hover::after{animation:shimmer 0.9s ease forwards}
@keyframes shimmer{100%{left:125%}}
.btn--full{width:100%}
.btn--sm{padding:6px 14px;font-size:0.7rem}
.btn--danger{border-color:var(--danger);color:var(--crimson-soft)}
.btn--danger:hover{background:rgba(122,43,43,0.2)}
.btn-icon{
  display:inline-flex;align-items:center;justify-content:center;
  width:28px;height:28px;border:1px solid transparent;background:transparent;
  color:var(--text-dim);cursor:pointer;transition:all var(--tr);border-radius:var(--radius);
  padding:0;
}
.btn-icon:hover{color:var(--gold-warm);border-color:var(--border-orn);background:var(--surface)}

/* ---- App Layout ---- */
.app-view{
  display:grid;
  grid-template-columns:260px 1fr 300px;
  grid-template-rows:52px 1fr;
  height:100vh;
}
.app-header{
  grid-column:1/-1;display:flex;align-items:center;justify-content:space-between;
  padding:0 24px;height:52px;background:var(--panel-dark);
  border-bottom:1px solid var(--border);position:relative;
}
.app-header::after{
  content:'';position:absolute;bottom:-1px;left:0;right:0;height:1px;
  background:linear-gradient(90deg,transparent,var(--gold-dim),transparent);opacity:0.3;
}
.header-brand{display:flex;align-items:center;gap:12px}
.header-crest{width:24px;height:24px;color:var(--gold-warm);opacity:0.8}
.header-title{font-family:var(--font-h);font-size:1.1rem;font-weight:700;letter-spacing:0.18em;color:var(--parchment)}
.header-sep{color:var(--text-dim);margin:0 4px}
.header-subtitle{font-family:var(--font-h);font-size:0.8rem;font-style:italic;color:var(--text-dim)}
.header-right{display:flex;align-items:center;gap:14px}
.header-user{font-family:var(--font-h);font-size:0.85rem;font-style:italic;color:var(--text-sec)}
.hamburger{display:none;background:none;border:none;color:var(--text-sec);cursor:pointer;padding:4px;font-size:1.2rem}
.panel-toggle{display:none;background:none;border:none;color:var(--text-sec);cursor:pointer;padding:4px;font-size:1rem}

/* ---- Conversation Sidebar ---- */
.conv-sidebar{
  background:linear-gradient(180deg,var(--panel-dark),var(--panel) 50%,var(--panel-dark));
  border-right:1px solid var(--border);display:flex;flex-direction:column;overflow:hidden;
}
.conv-sidebar-header{padding:16px;flex-shrink:0}
.conv-list{flex:1;overflow-y:auto;padding:0 8px 16px}
.conv-accordion{margin-top:10px;border-top:1px solid rgba(200,184,152,0.08)}
.conv-accordion-btn{
  width:100%;display:flex;align-items:center;justify-content:space-between;
  padding:12px 12px 8px;background:none;border:none;cursor:pointer;
  color:var(--text-sec);font-family:var(--font-ui);font-size:0.7rem;font-weight:700;
  letter-spacing:0.12em;text-transform:uppercase;text-align:left;
}
.conv-accordion-btn:hover{color:var(--gold-warm)}
.conv-accordion-icon{color:var(--text-dim);font-size:0.82rem}
.conv-accordion-body{padding-bottom:4px}
.conv-accordion-body.hidden{display:none}
.conv-group-label{
  font-family:var(--font-ui);font-size:0.6rem;font-weight:700;text-transform:uppercase;
  letter-spacing:0.14em;color:var(--text-dim);padding:14px 12px 6px;
}
.conv-item{
  display:flex;align-items:center;gap:6px;padding:9px 12px;margin-bottom:2px;
  cursor:pointer;transition:all var(--tr);border-radius:var(--radius);
  border-left:2px solid transparent;position:relative;
}
.conv-item:hover{background:var(--surface);color:var(--text)}
.conv-item.active{color:var(--gold-warm);background:rgba(212,168,67,0.06);border-left-color:var(--gold-warm)}
.conv-item.active .conv-title{color:var(--gold-warm)}
.conv-title-wrap{display:flex;flex-direction:column;gap:2px;flex:1;min-width:0}
.conv-title{
  font-family:var(--font-h);font-size:0.9rem;color:var(--text-sec);
  overflow:hidden;text-overflow:ellipsis;white-space:nowrap;
}
.conv-project{
  font-family:var(--font-ui);font-size:0.62rem;letter-spacing:0.06em;
  color:var(--text-dim);text-transform:uppercase;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;
}
.conv-actions{display:none;gap:2px;flex-shrink:0}
.conv-item:hover .conv-actions{display:flex}
.conv-rename-input{
  flex:1;padding:2px 6px;background:var(--panel-dark);border:1px solid var(--gold-dim);
  color:var(--text);font-family:var(--font-h);font-size:0.9rem;outline:none;
  border-radius:var(--radius);
}
.conv-empty{padding:24px 16px;text-align:center;color:var(--text-dim);font-style:italic;font-size:0.85rem}

/* ---- Chat Area ---- */
.chat-main{display:flex;flex-direction:column;overflow:hidden;
  background:radial-gradient(ellipse 600px 400px at 30% 30%,rgba(212,168,67,0.02),transparent),var(--bg);
}
.chat-messages{
  flex:1;overflow-y:auto;padding:24px 28px;scroll-behavior:smooth;
  background:repeating-linear-gradient(180deg,transparent,transparent 3px,rgba(200,184,152,0.008) 3px,rgba(200,184,152,0.008) 4px);
}
.chat-empty{
  display:flex;flex-direction:column;align-items:center;justify-content:center;
  height:100%;color:var(--text-dim);text-align:center;gap:12px;
}
.chat-empty-icon{font-size:3rem;opacity:0.3}
.chat-empty-title{font-family:var(--font-h);font-size:1.3rem;color:var(--text-sec)}
.chat-empty-hint{font-size:0.85rem;font-style:italic}
.chat-message{display:flex;gap:14px;margin-bottom:24px;max-width:82%}
.chat-message--user{margin-left:auto;flex-direction:row-reverse}
.chat-avatar{
  flex-shrink:0;width:34px;height:34px;display:flex;align-items:center;justify-content:center;
  font-size:0.65rem;font-weight:700;text-transform:uppercase;letter-spacing:0.04em;
  border-radius:var(--radius);
}
.chat-message--user .chat-avatar{
  background:var(--surface);color:var(--text-sec);border:1px solid var(--border);font-family:var(--font-ui);
}
.chat-message--merlin .chat-avatar{
  background:linear-gradient(135deg,var(--gold-dim),var(--gold-warm));color:var(--bg);
  font-family:var(--font-h);font-size:1rem;font-weight:700;
  box-shadow:0 0 12px rgba(212,168,67,0.15);
}
.chat-body{flex:1;min-width:0}
.chat-bubble{
  padding:14px 18px;line-height:1.65;white-space:pre-wrap;word-break:break-word;
  font-size:0.92rem;border-radius:var(--radius);
}
.chat-message--user .chat-bubble{
  background:var(--surface);border:1px solid var(--border);border-left:3px solid var(--gold-dim);
}
.chat-message--merlin .chat-bubble{
  background:var(--panel);border:1px solid var(--border-sub);
  border-left:3px solid rgba(212,168,67,0.2);color:var(--parchment);
  box-shadow:0 0 25px rgba(224,194,106,0.10);
}
.chat-bubble--error{border-color:var(--crimson)!important;border-left-color:var(--crimson)!important;color:var(--crimson-soft)}

.chat-sources{margin-top:10px;border:1px solid var(--border-sub);border-top:1px solid var(--border);padding:10px 14px;background:rgba(20,18,16,0.6);border-radius:var(--radius)}
.chat-sources summary{font-family:var(--font-h);font-size:0.8rem;font-style:italic;color:var(--text-dim);cursor:pointer;user-select:none}
.chat-sources summary:hover{color:var(--text-sec)}
.chat-sources-list{display:flex;flex-wrap:wrap;gap:6px;margin-top:10px}
.source-pill{
  font-family:var(--font-ui);font-size:0.68rem;padding:4px 10px;
  border:1px solid var(--border);background:var(--panel-dark);color:var(--text-sec);border-radius:var(--radius);
}
.source-pill strong{color:var(--parchment)}
.source-pill a{color:var(--gold-warm);text-decoration:none}
.source-pill a:hover{color:var(--gold-bright);text-decoration:underline}
.source-score{color:var(--gold-warm);margin-left:4px}
.source-pill--web{border-color:var(--verdigris);background:rgba(79,143,107,0.10)}
.source-pill--web strong{color:var(--verdigris)}

/* ---- Performance Indicator ---- */
.chat-perf{
  display:flex;gap:10px;margin-top:6px;font-family:var(--font-mono);font-size:0.65rem;color:var(--text-dim);
}
.chat-perf span{display:inline-flex;align-items:center;gap:3px}
.perf-value{color:var(--gold-warm);font-weight:600}

/* ---- Export Dropdown ---- */
.export-dropdown{
  position:absolute;z-index:50;background:var(--panel-dark);border:1px solid var(--border-orn);
  padding:6px 0;min-width:120px;box-shadow:var(--shadow-sm);border-radius:var(--radius);
}
.export-dropdown a{
  display:block;padding:6px 14px;font-family:var(--font-ui);font-size:0.75rem;
  color:var(--text);text-decoration:none;transition:background var(--tr);
}
.export-dropdown a:hover{background:var(--surface);color:var(--gold-warm)}

/* ---- Generate Modal ---- */
.generate-overlay{
  position:fixed;top:0;left:0;right:0;bottom:0;background:rgba(0,0,0,0.6);
  z-index:200;display:flex;align-items:center;justify-content:center;
}
.generate-modal{
  background:var(--panel-dark);border:1px solid var(--border-orn);
  padding:28px;width:90%;max-width:480px;box-shadow:var(--shadow);border-radius:var(--radius);
}
.generate-modal h3{font-family:var(--font-h);font-size:1.1rem;color:var(--parchment);margin-bottom:16px}
.generate-modal textarea{
  width:100%;min-height:80px;padding:10px;background:var(--panel);
  border:1px solid var(--border);color:var(--text);font-family:var(--font-h);
  font-size:0.9rem;resize:vertical;outline:none;margin-bottom:12px;border-radius:var(--radius);
}
.generate-modal textarea:focus{border-color:var(--gold-dim)}
.generate-modal select{
  width:100%;padding:8px;background:var(--panel);border:1px solid var(--border);
  color:var(--text);font-family:var(--font-h);font-size:0.85rem;margin-bottom:16px;
  outline:none;border-radius:var(--radius);-webkit-appearance:none;
}
.generate-modal select:focus{border-color:var(--gold-dim)}
.generate-modal .modal-actions{display:flex;gap:10px;justify-content:flex-end}

@keyframes pulse{0%,100%{opacity:1}50%{opacity:0.5}}
.thinking{animation:pulse 1.5s ease-in-out infinite}

/* ---- Composer ---- */
.chat-composer{
  padding:16px 28px;border-top:1px solid var(--border);background:var(--panel-dark);position:relative;
}
.chat-composer::before{
  content:'';position:absolute;top:-1px;left:20%;right:20%;height:1px;
  background:linear-gradient(90deg,transparent,var(--gold-dim),transparent);opacity:0.3;
}
.chat-input-wrap{display:flex;gap:12px;align-items:flex-end}
.chat-textarea{
  flex:1;min-height:46px;max-height:160px;padding:11px 14px;
  background:var(--panel);border:1px solid var(--border);color:var(--text);
  font-family:var(--font-h);font-size:0.95rem;line-height:1.5;resize:none;outline:none;
  transition:border-color var(--tr);border-radius:var(--radius);
}
.chat-textarea:focus{border-color:var(--gold-dim);box-shadow:0 0 0 3px rgba(224,194,106,0.08)}
.chat-textarea::placeholder{color:var(--text-dim);font-style:italic}
.chat-actions{display:flex;gap:8px;margin-top:8px;align-items:center}
.chat-hint{font-family:var(--font-ui);font-size:0.65rem;color:var(--text-dim);margin-left:auto}
.chat-status{font-family:var(--font-h);font-size:0.8rem;font-style:italic;color:var(--gold-warm)}

/* ---- Right Panel ---- */
.right-panel{
  background:linear-gradient(180deg,var(--panel-dark),var(--panel) 50%,var(--panel-dark));
  border-left:1px solid var(--border);display:flex;flex-direction:column;overflow-y:auto;
}
.panel-section{padding:16px}
.panel-title{
  font-family:var(--font-h);font-size:0.8rem;font-style:italic;color:var(--text-dim);
  margin-bottom:12px;text-align:center;
}
.panel-divider{height:1px;margin:0 16px;background:linear-gradient(90deg,transparent,var(--border),transparent)}
.settings-row{display:flex;gap:8px;margin-bottom:10px}
.settings-row:last-child{margin-bottom:0}
.settings-field{flex:1}
.settings-label{
  display:block;font-family:var(--font-ui);font-size:0.6rem;font-weight:600;
  text-transform:uppercase;letter-spacing:0.08em;color:var(--text-dim);margin-bottom:4px;
}
.settings-input,.settings-select{
  width:100%;padding:6px 10px;background:var(--panel-dark);
  border:1px solid var(--border-sub);border-bottom:1px solid var(--border);
  color:var(--text);font-family:var(--font-h);font-size:0.8rem;outline:none;
  border-radius:var(--radius);-webkit-appearance:none;
}
.settings-input:focus,.settings-select:focus{border-color:var(--gold-dim);box-shadow:0 0 0 3px rgba(224,194,106,0.08)}

.panel-actions{display:flex;gap:8px;margin-bottom:12px}
.drop-zone{
  border:2px dashed var(--border-orn);padding:20px 12px;text-align:center;
  background:rgba(30,27,24,0.5);transition:all var(--tr);cursor:pointer;
  border-radius:var(--radius);font-size:0.8rem;color:var(--text-dim);font-style:italic;
}
.drop-zone.active{border-color:var(--gold-warm);background:rgba(212,168,67,0.04)}
.file-label{color:var(--gold-warm);cursor:pointer;text-decoration:underline}
.upload-status{font-family:var(--font-h);font-size:0.8rem;font-style:italic;color:var(--text-sec);margin-top:8px}
.upload-status--ok{color:var(--success)}
.upload-status--err{color:var(--crimson-soft)}
.file-list{margin-top:12px;max-height:300px;overflow-y:auto}
.file-section-title{
  font-family:var(--font-ui);font-size:0.6rem;font-weight:700;text-transform:uppercase;
  letter-spacing:0.1em;color:var(--text-dim);margin:10px 0 6px;padding-bottom:4px;
  border-bottom:1px solid var(--border-sub);
}
.file-item{
  display:flex;justify-content:space-between;align-items:center;padding:5px 0;
  font-size:0.78rem;border-bottom:1px solid var(--border-sub);
}
.file-item:last-child{border-bottom:none}
.file-name{color:var(--text);font-weight:600;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;flex:1}
.file-size{color:var(--text-dim);font-family:var(--font-mono);font-size:0.68rem;margin-left:8px}

.health-bar{
  padding:12px 16px;margin-top:auto;border-top:1px solid var(--border-sub);
  font-family:var(--font-ui);font-size:0.65rem;color:var(--text-dim);text-align:center;
}

/* ---- Scrollbar ---- */
::-webkit-scrollbar{width:6px}
::-webkit-scrollbar-track{background:transparent}
::-webkit-scrollbar-thumb{background:var(--scroll);border-radius:0}
::-webkit-scrollbar-thumb:hover{background:var(--hover)}

/* ---- Responsive ---- */
@media(max-width:1024px){
  .app-view{grid-template-columns:220px 1fr 260px}
}
@media(max-width:768px){
  .app-view{grid-template-columns:1fr;grid-template-rows:52px 1fr}
  .conv-sidebar{display:none;position:fixed;top:52px;left:0;bottom:0;width:280px;z-index:100;
    box-shadow:4px 0 20px rgba(0,0,0,0.5)}
  .conv-sidebar.open{display:flex}
  .right-panel{display:none;position:fixed;top:52px;right:0;bottom:0;width:300px;z-index:100;
    box-shadow:-4px 0 20px rgba(0,0,0,0.5)}
  .right-panel.open{display:flex}
  .hamburger{display:block}
  .panel-toggle{display:block}
  .chat-message{max-width:95%}
  .header-subtitle{display:none}
}
</style>
</head>
<body>

<!-- ==================== LOGIN VIEW ==================== -->
<div id="login-view" class="login-view">
  <div class="login-card">
    <div class="corner-bl"></div>
    <div class="corner-br"></div>
    <div class="login-brand">
      <svg class="login-crest" viewBox="0 0 56 56" fill="none">
        <path d="M28 6l5 14h14l-11 8 4 14-12-9-12 9 4-14L9 20h14z" stroke="currentColor" stroke-width="1.5" fill="currentColor" fill-opacity="0.12"/>
        <line x1="28" y1="38" x2="28" y2="52" stroke="currentColor" stroke-width="2.5" stroke-linecap="round"/>
      </svg>
      <div class="login-wordmark">MERLIN</div>
      <div class="login-accent"></div>
      <div class="login-tagline">Knowledge Retrieval System</div>
    </div>

    <div class="auth-tabs">
      <button class="auth-tab active" data-tab="login">Login</button>
      <button class="auth-tab" data-tab="register">Register</button>
    </div>

    <form id="login-form">
      <div class="form-group">
        <label class="form-label">Username</label>
        <input class="form-input" type="text" id="login-username" placeholder="Enter username" autocomplete="username"/>
      </div>
      <div class="form-group">
        <label class="form-label">Password</label>
        <input class="form-input" type="password" id="login-password" placeholder="Enter password" autocomplete="current-password"/>
      </div>
      <div class="login-error" id="login-error"></div>
      <button type="submit" class="btn btn--primary btn--full">Enter the Archive</button>
    </form>

    <form id="register-form" class="hidden">
      <div class="form-group">
        <label class="form-label">Username</label>
        <input class="form-input" type="text" id="reg-username" placeholder="Choose a username (4+ chars)" autocomplete="username"/>
      </div>
      <div class="form-group">
        <label class="form-label">Password</label>
        <input class="form-input" type="password" id="reg-password" placeholder="Choose a password (8+ chars)" autocomplete="new-password"/>
      </div>
      <div class="form-group">
        <label class="form-label">Confirm Password</label>
        <input class="form-input" type="password" id="reg-confirm" placeholder="Confirm password" autocomplete="new-password"/>
      </div>
      <div class="login-error" id="register-error"></div>
      <button type="submit" class="btn btn--primary btn--full">Create Account</button>
    </form>
  </div>
</div>

<!-- ==================== APP VIEW ==================== -->
<div id="app-view" class="app-view hidden">

  <!-- Header -->
  <header class="app-header">
    <div class="header-brand">
      <button class="hamburger" id="hamburger-btn" title="Toggle sidebar">&#9776;</button>
      <svg class="header-crest" viewBox="0 0 24 24" fill="none">
        <path d="M12 2l2.2 6.8H21l-5.5 4 2.1 6.5L12 15l-5.6 4.3 2.1-6.5L3 8.8h6.8z" stroke="currentColor" stroke-width="1.2" fill="currentColor" fill-opacity="0.15"/>
        <line x1="12" y1="17" x2="12" y2="22" stroke="currentColor" stroke-width="1.8" stroke-linecap="round"/>
      </svg>
      <span class="header-title">MERLIN</span>
      <span class="header-sep">&middot;</span>
      <span class="header-subtitle">Knowledge Retrieval System</span>
    </div>
    <div class="header-right">
      <span class="header-user" id="display-username"></span>
      <button class="panel-toggle" id="panel-toggle-btn" title="Toggle settings">&#9881;</button>
      <button class="btn btn--sm" id="logout-btn">Logout</button>
    </div>
  </header>

  <!-- Conversation Sidebar -->
  <aside class="conv-sidebar" id="conv-sidebar">
    <div class="conv-sidebar-header">
      <button class="btn btn--primary btn--full" id="new-conv-btn">+ New Conversation</button>
    </div>
    <div class="conv-list" id="conv-list">
      <div class="conv-empty">No conversations yet</div>
    </div>
  </aside>

  <!-- Chat Area -->
  <main class="chat-main">
    <div class="chat-messages" id="chat-messages">
      <div class="chat-empty" id="chat-empty">
        <div class="chat-empty-icon">&#9733;</div>
        <div class="chat-empty-title">Welcome to Merlin</div>
        <div class="chat-empty-hint">Start a conversation or select one from the sidebar</div>
      </div>
    </div>
    <div class="chat-composer">
      <div class="chat-input-wrap">
        <textarea class="chat-textarea" id="chat-input" placeholder="Ask about your documents..." rows="1"></textarea>
        <button class="btn btn--primary" id="send-btn">Send</button>
        <button class="btn btn--sm" id="generate-btn" title="Generate a file">&#128196; Generate</button>
      </div>
      <div class="chat-actions">
        <span class="chat-status" id="chat-status"></span>
        <span class="chat-hint">Enter to send &middot; Shift+Enter for newline</span>
      </div>
    </div>
  </main>

  <!-- Right Panel -->
  <aside class="right-panel" id="right-panel">
    <div class="panel-section">
      <div class="panel-title">Query Settings</div>
      <div class="settings-row">
        <div class="settings-field">
          <label class="settings-label">Mode</label>
          <select class="settings-select" id="mode-select">
            <option value="research" selected>Research (docs + web)</option>
            <option value="grounded">Grounded (Document Analysis only)</option>
            <option value="general">General (LLM only)</option>
          </select>
        </div>
      </div>
      <div class="settings-row">
        <div class="settings-field">
          <label class="settings-label">Top K</label>
          <input class="settings-input" type="number" id="topk-input" value="12" min="1" max="20"/>
        </div>
        <div class="settings-field">
          <label class="settings-label">History</label>
          <input class="settings-input" type="number" id="hist-input" value="12" min="0" max="50"/>
        </div>
      </div>
      <div class="settings-row">
        <div class="settings-field">
          <label class="settings-label">Web Results</label>
          <input class="settings-input" type="number" id="web-results-input" value="5" min="1" max="8"/>
        </div>
      </div>
    </div>

    <div class="panel-divider"></div>

    <div class="panel-section">
      <div class="panel-title">Projects</div>
      <div class="settings-row">
        <div class="settings-field">
          <label class="settings-label">Active Project</label>
          <select class="settings-select" id="project-select">
            <option value="">No project selected</option>
          </select>
        </div>
      </div>
      <div class="settings-row">
        <div class="settings-field">
          <label class="settings-label">New Project</label>
          <input class="settings-input" type="text" id="project-name-input" placeholder="Project name"/>
        </div>
      </div>
      <div class="panel-actions">
        <button class="btn btn--sm btn--primary" id="create-project-btn">Create</button>
        <button class="btn btn--sm" id="delete-project-btn">Delete</button>
      </div>
      <div class="upload-status" id="project-status"></div>
    </div>

    <div class="panel-divider"></div>

    <div class="panel-section">
      <div class="panel-title">Documents</div>
      <div class="panel-actions">
        <button class="btn btn--sm btn--primary" id="ingest-btn">Re-Ingest</button>
        <button class="btn btn--sm" id="refresh-files-btn">Refresh</button>
      </div>
      <div class="drop-zone" id="drop-zone">
        Drop files here or <label for="file-input" class="file-label">browse</label>
        <input type="file" id="file-input" multiple style="display:none"/>
      </div>
      <div class="upload-status" id="upload-status"></div>
      <div class="file-list" id="file-list"></div>
    </div>

    <div class="health-bar" id="health-bar">checking&hellip;</div>
  </aside>
</div>

<!-- Generate File Modal -->
<div class="generate-overlay hidden" id="generate-overlay">
  <div class="generate-modal">
    <h3>Generate File</h3>
    <textarea id="gen-prompt" placeholder="Describe what you want to generate..."></textarea>
    <select id="gen-format">
      <option value="txt">Text (.txt)</option>
      <option value="csv">Spreadsheet (.csv)</option>
      <option value="docx">Document (.docx)</option>
      <option value="png">Chart (.png)</option>
    </select>
    <div class="modal-actions">
      <button class="btn btn--sm" id="gen-cancel">Cancel</button>
      <button class="btn btn--sm btn--primary" id="gen-submit">Generate</button>
    </div>
  </div>
</div>

<script>
/* ============================================================
   MERLIN RAG — Client SPA
   ============================================================ */
(function(){
'use strict';

// ---- State ----
const S = {
  user: null,
  projects: [],
  activeProject: null,
  conversations: [],
  activeConvo: null,
  convoSections: { project: true, other: true },
};

// ---- DOM helpers ----
const $ = id => document.getElementById(id);
function esc(s) {
  return String(s).replace(/[&<>"']/g, c => ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'})[c] || c);
}
function fmtBytes(n) {
  if (n < 1024) return n + ' B';
  const k = n / 1024;
  if (k < 1024) return k.toFixed(1) + ' KB';
  const m = k / 1024;
  if (m < 1024) return m.toFixed(1) + ' MB';
  return (m / 1024).toFixed(2) + ' GB';
}

// ---- API ----
async function api(method, url, body) {
  const opts = { method, headers: {} };
  if (body !== undefined) {
    opts.headers['Content-Type'] = 'application/json';
    opts.body = JSON.stringify(body);
  }
  const r = await fetch(url, opts);
  if (r.status === 401) { showLogin(); throw new Error('Session expired'); }
  const text = await r.text();
  let data;
  try { data = JSON.parse(text); } catch { throw new Error(text); }
  if (!r.ok) throw new Error(data.detail || JSON.stringify(data));
  return data;
}
const GET = url => api('GET', url);
const POST = (url, body) => api('POST', url, body);
const PATCH = (url, body) => api('PATCH', url, body);
const DEL = url => api('DELETE', url);

// ---- Views ----
function showLogin() {
  S.user = null;
  S.projects = [];
  S.activeProject = null;
  $('login-view').classList.remove('hidden');
  $('app-view').classList.add('hidden');
}
function showApp() {
  $('login-view').classList.add('hidden');
  $('app-view').classList.remove('hidden');
  $('display-username').textContent = S.user ? S.user.username : '';
}

// ---- Auth ----
async function checkSession() {
  try {
    const data = await GET('/auth/me');
    S.user = data;
    showApp();
    await loadProjects();
    await loadConversations();
    await loadFiles();
    checkHealth();
  } catch {
    showLogin();
  }
}

async function login(username, password) {
  const data = await POST('/auth/login', { username, password });
  S.user = data;
  showApp();
  await loadProjects();
  await loadConversations();
  await loadFiles();
  checkHealth();
}

async function register(username, password) {
  const data = await POST('/auth/register', { username, password });
  S.user = data;
  showApp();
  await loadProjects();
  await loadConversations();
  await loadFiles();
  checkHealth();
}

async function logout() {
  try { await POST('/auth/logout'); } catch {}
  S.user = null;
  S.projects = [];
  S.activeProject = null;
  S.conversations = [];
  S.activeConvo = null;
  showLogin();
}

// ---- Projects ----
function renderProjects() {
  const select = $('project-select');
  if (!select) return;
  let html = '<option value="">No project selected</option>';
  S.projects.forEach(function(p) {
    const selected = p.id === S.activeProject ? ' selected' : '';
    html += '<option value="' + esc(p.id) + '"' + selected + '>' + esc(p.name) + ' (' + p.file_count + ')</option>';
  });
  select.innerHTML = html;
  $('project-status').textContent = S.activeProject ? '' : 'Choose a project to scope uploads and retrieval.';
}

async function loadProjects() {
  try {
    S.projects = await GET('/projects');
    if (S.activeProject && !S.projects.find(function(p) { return p.id === S.activeProject; })) {
      S.activeProject = null;
    }
    renderProjects();
  } catch (e) {
    console.error('Failed to load projects:', e);
  }
}

async function createProject() {
  const input = $('project-name-input');
  const name = input.value.trim();
  if (!name) {
    $('project-status').textContent = 'Enter a project name first.';
    return;
  }
  try {
    const project = await POST('/projects', { name: name });
    input.value = '';
    S.activeProject = project.id;
    await loadProjects();
    await loadFiles();
    $('project-status').textContent = 'Project created.';
  } catch (e) {
    $('project-status').textContent = 'Project error: ' + e.message;
  }
}

async function deleteActiveProject() {
  if (!S.activeProject) {
    $('project-status').textContent = 'Select a project first.';
    return;
  }
  const project = S.projects.find(function(p) { return p.id === S.activeProject; });
  const label = project ? project.name : 'this project';
  if (!confirm('Delete ' + label + '? Its files will be removed from retrieval.')) return;
  try {
    await DEL('/projects/' + encodeURIComponent(S.activeProject));
    const deletedId = S.activeProject;
    S.activeProject = null;
    S.conversations = S.conversations.map(function(c) {
      return c.project_id === deletedId ? Object.assign({}, c, { project_id: null, project_name: null }) : c;
    });
    renderConversations();
    await loadProjects();
    await loadFiles();
    $('project-status').textContent = 'Project deleted.';
  } catch (e) {
    $('project-status').textContent = 'Delete failed: ' + e.message;
  }
}

async function setActiveProject(projectId, syncConversation) {
  S.activeProject = projectId || null;
  renderProjects();
  await loadFiles();
  if (syncConversation && S.activeConvo) {
    try {
      await PATCH('/conversations/' + encodeURIComponent(S.activeConvo) + '/project', { project_id: S.activeProject });
      S.conversations = S.conversations.map(function(c) {
        if (c.id !== S.activeConvo) return c;
        const project = S.projects.find(function(p) { return p.id === S.activeProject; });
        return Object.assign({}, c, {
          project_id: S.activeProject,
          project_name: project ? project.name : null,
        });
      });
      renderConversations();
    } catch (e) {
      $('project-status').textContent = 'Project assignment failed: ' + e.message;
    }
  }
}

// ---- Conversations ----
function groupConvos(convos) {
  const d = new Date(); d.setHours(0,0,0,0);
  const todayTs = Math.floor(d.getTime() / 1000);
  const weekTs = todayTs - 7 * 86400;
  const today = [], week = [], older = [];
  for (const c of convos) {
    if (c.updated_at >= todayTs) today.push(c);
    else if (c.updated_at >= weekTs) week.push(c);
    else older.push(c);
  }
  return [
    { label: 'Today', items: today },
    { label: 'Previous 7 Days', items: week },
    { label: 'Older', items: older },
  ].filter(g => g.items.length > 0);
}

function renderConversationItems(convos) {
  const groups = groupConvos(convos);
  let html = '';
  for (const g of groups) {
    html += '<div class="conv-group-label">' + esc(g.label) + '</div>';
    for (const c of g.items) {
      const active = c.id === S.activeConvo ? ' active' : '';
      const projectMeta = c.project_name ? '<span class="conv-project">' + esc(c.project_name) + '</span>' : '';
      html += '<div class="conv-item' + active + '" data-conv="' + esc(c.id) + '">'
        + '<span class="conv-title-wrap"><span class="conv-title">' + esc(c.title) + '</span>' + projectMeta + '</span>'
        + '<span class="conv-actions">'
        + '<button class="btn-icon conv-export" title="Export">&#128190;</button>'
        + '<button class="btn-icon conv-rename" title="Rename">&#9998;</button>'
        + '<button class="btn-icon conv-delete" title="Delete">&#128465;</button>'
        + '</span></div>';
    }
  }
  return html;
}

function renderConversationAccordion(key, label, convos, emptyLabel) {
  const open = !!S.convoSections[key];
  const icon = open ? '&#9662;' : '&#9656;';
  const bodyClass = open ? 'conv-accordion-body' : 'conv-accordion-body hidden';
  const body = convos.length
    ? renderConversationItems(convos)
    : '<div class="conv-empty">' + esc(emptyLabel) + '</div>';
  return '<div class="conv-accordion" data-accordion="' + esc(key) + '">'
    + '<button class="conv-accordion-btn" data-toggle-accordion="' + esc(key) + '">'
    + '<span>' + esc(label) + ' (' + convos.length + ')</span>'
    + '<span class="conv-accordion-icon">' + icon + '</span>'
    + '</button>'
    + '<div class="' + bodyClass + '">' + body + '</div></div>';
}

function renderConversations() {
  const el = $('conv-list');
  if (!S.conversations.length) {
    el.innerHTML = '<div class="conv-empty">No conversations yet</div>';
    return;
  }
  let html = '';
  if (S.activeProject) {
    const activeProject = S.projects.find(function(p) { return p.id === S.activeProject; });
    const activeName = activeProject ? activeProject.name : 'Selected Project';
    const projectConvos = S.conversations.filter(function(c) { return c.project_id === S.activeProject; });
    const otherConvos = S.conversations.filter(function(c) { return c.project_id !== S.activeProject; });
    html += renderConversationAccordion(
      'project',
      activeName + ' Chats',
      projectConvos,
      'No conversations in this project yet'
    );
    html += renderConversationAccordion(
      'other',
      'Other Chats',
      otherConvos,
      'No other conversations'
    );
  } else {
    html = renderConversationItems(S.conversations);
  }
  el.innerHTML = html;
}

async function loadConversations() {
  try {
    S.conversations = await GET('/conversations');
    renderConversations();
  } catch (e) {
    console.error('Failed to load conversations:', e);
  }
}

async function createConversation() {
  const conv = await POST('/conversations', { project_id: S.activeProject });
  S.conversations.unshift(conv);
  S.activeConvo = conv.id;
  renderConversations();
  clearChatView();
  $('chat-input').focus();
}

async function selectConversation(id) {
  if (S.activeConvo === id) return;
  S.activeConvo = id;
  const convo = S.conversations.find(function(c) { return c.id === id; });
  S.activeProject = convo ? (convo.project_id || null) : null;
  renderProjects();
  renderConversations();
  await loadFiles();
  await loadMessages(id);
}

async function loadMessages(convoId) {
  const msgsEl = $('chat-messages');
  msgsEl.innerHTML = '';
  try {
    const res = await GET('/conversations/' + convoId + '/messages');
    const msgs = res.messages || [];
    if (typeof res.conversation !== 'undefined') {
      S.activeProject = res.conversation.project_id || null;
      renderProjects();
      await loadFiles();
    }
    if (!msgs.length) {
      msgsEl.innerHTML = '<div class="chat-empty"><div class="chat-empty-icon">&#9733;</div>'
        + '<div class="chat-empty-title">New Conversation</div>'
        + '<div class="chat-empty-hint">Ask a question to begin</div></div>';
      return;
    }
    for (const m of msgs) {
      renderMessage(m.role, m.content);
    }
    msgsEl.scrollTop = msgsEl.scrollHeight;
  } catch (e) {
    msgsEl.innerHTML = '<div class="chat-empty"><div class="chat-empty-hint">Error loading messages: ' + esc(e.message) + '</div></div>';
  }
}

function startRename(convId) {
  const item = document.querySelector('[data-conv="' + convId + '"]');
  if (!item) return;
  const titleEl = item.querySelector('.conv-title');
  const original = titleEl.textContent;
  const input = document.createElement('input');
  input.className = 'conv-rename-input';
  input.value = original;
  titleEl.replaceWith(input);
  input.focus();
  input.select();

  let saved = false;
  async function save() {
    if (saved) return;
    saved = true;
    const title = input.value.trim();
    if (title && title !== original) {
      try { await PATCH('/conversations/' + convId, { title }); } catch {}
    }
    await loadConversations();
  }
  input.addEventListener('keydown', e => {
    if (e.key === 'Enter') { e.preventDefault(); save(); }
    if (e.key === 'Escape') { saved = true; loadConversations(); }
  });
  input.addEventListener('blur', save);
}

async function deleteConversation(convId) {
  if (!confirm('Delete this conversation? This cannot be undone.')) return;
  try {
    await DEL('/conversations/' + convId);
    if (S.activeConvo === convId) {
      S.activeConvo = null;
      clearChatView();
    }
    await loadConversations();
  } catch (e) {
    alert('Delete failed: ' + e.message);
  }
}

function clearChatView() {
  $('chat-messages').innerHTML = '<div class="chat-empty"><div class="chat-empty-icon">&#9733;</div>'
    + '<div class="chat-empty-title">Welcome to Merlin</div>'
    + '<div class="chat-empty-hint">Start a conversation or select one from the sidebar</div></div>';
  $('chat-status').textContent = '';
}

// ---- Chat ----
function renderMessage(role, content, sources, performance) {
  const msgsEl = $('chat-messages');
  // Remove empty state if present
  const empty = msgsEl.querySelector('.chat-empty');
  if (empty) empty.remove();

  const isUser = role === 'user';
  const wrap = document.createElement('div');
  wrap.className = 'chat-message ' + (isUser ? 'chat-message--user' : 'chat-message--merlin');

  const avatar = document.createElement('div');
  avatar.className = 'chat-avatar';
  avatar.textContent = isUser ? 'YOU' : 'M';

  const body = document.createElement('div');
  body.className = 'chat-body';

  const bubble = document.createElement('div');
  bubble.className = 'chat-bubble';
  if (role === 'error') bubble.classList.add('chat-bubble--error');
  bubble.textContent = content;

  body.appendChild(bubble);

  if (sources && sources.length) {
    const det = document.createElement('details');
    det.className = 'chat-sources';
    const sum = document.createElement('summary');
    sum.textContent = 'Sources (' + sources.length + ')';
    det.appendChild(sum);
    const list = document.createElement('div');
    list.className = 'chat-sources-list';
    sources.forEach(function(s) {
      const pill = document.createElement('span');
      if (s.url) {
        pill.className = 'source-pill source-pill--web';
        pill.innerHTML = '&#128279; <a href="' + esc(s.url) + '" target="_blank" rel="noopener">'
          + '<strong>' + esc(s.source) + '</strong></a>';
      } else {
        pill.className = 'source-pill';
        pill.innerHTML = '<strong>' + esc(s.source) + '</strong> #' + s.chunk_index
          + '<span class="source-score">' + (s.score ? s.score.toFixed(3) : '') + '</span>';
      }
      list.appendChild(pill);
    });
    det.appendChild(list);
    body.appendChild(det);
  }

  // Performance indicator for assistant messages
  if (performance && performance.tokens_per_sec && role === 'assistant') {
    const perf = document.createElement('div');
    perf.className = 'chat-perf';
    perf.innerHTML = '<span><span class="perf-value">' + performance.tokens_per_sec + '</span> tok/s</span>'
      + '<span><span class="perf-value">' + performance.tokens + '</span> tokens</span>'
      + '<span><span class="perf-value">' + performance.total_duration + '</span>s total</span>';
    body.appendChild(perf);
  }

  wrap.appendChild(avatar);
  wrap.appendChild(body);
  msgsEl.appendChild(wrap);
  msgsEl.scrollTop = msgsEl.scrollHeight;
}

function showThinking() {
  const msgsEl = $('chat-messages');
  const empty = msgsEl.querySelector('.chat-empty');
  if (empty) empty.remove();

  const wrap = document.createElement('div');
  wrap.id = 'thinking-msg';
  wrap.className = 'chat-message chat-message--merlin';
  wrap.innerHTML = '<div class="chat-avatar">M</div>'
    + '<div class="chat-body"><div class="chat-bubble thinking">Consulting the archives...</div></div>';
  msgsEl.appendChild(wrap);
  msgsEl.scrollTop = msgsEl.scrollHeight;
}

function hideThinking() {
  const el = $('thinking-msg');
  if (el) el.remove();
}

async function sendMessage() {
  const input = $('chat-input');
  const question = input.value.trim();
  if (!question) return;

  // Auto-create conversation if none selected
  if (!S.activeConvo) {
    try {
      const conv = await POST('/conversations', { project_id: S.activeProject });
      S.activeConvo = conv.id;
      S.conversations.unshift(conv);
      renderConversations();
    } catch (e) {
      $('chat-status').textContent = 'Error: ' + e.message;
      return;
    }
  }

  renderMessage('user', question);
  input.value = '';
  input.style.height = 'auto';
  $('chat-status').textContent = 'Thinking...';
  showThinking();

  const mode = $('mode-select').value;
  const topK = parseInt($('topk-input').value, 10) || 12;
  const histLimit = parseInt($('hist-input').value, 10) || 12;
  const webSearch = mode === 'research';
  const webResults = parseInt($('web-results-input').value, 10) || 5;

  try {
    const res = await POST('/chat', {
      convo_id: S.activeConvo,
      question: question,
      project_id: S.activeProject,
      mode: mode,
      top_k: topK,
      history_limit: histLimit,
      web_search: webSearch,
      web_results: webResults,
    });
    hideThinking();
    renderMessage('assistant', res.answer, res.sources, res.performance);
    $('chat-status').textContent = '';
    // Refresh conversations (title may have updated)
    await loadConversations();
  } catch (e) {
    hideThinking();
    renderMessage('error', 'Error: ' + e.message);
    $('chat-status').textContent = '';
  }
}

// ---- Files ----
async function loadFiles() {
  const el = $('file-list');
  if (!S.activeProject) {
    el.innerHTML = '<div style="color:var(--text-dim);font-size:0.78rem;font-style:italic;padding:8px 0">Select a project to view its files.</div>';
    return;
  }
  el.innerHTML = '<div style="color:var(--text-dim);font-size:0.78rem;font-style:italic;padding:8px 0">Loading...</div>';
  try {
    const res = await GET('/files?project_id=' + encodeURIComponent(S.activeProject));
    const text = res.text || [];
    let html = '';
    html += '<div class="file-section-title">Project Files &middot; ' + text.length + ' file(s)</div>';
    if (text.length) {
      text.forEach(function(f) {
        html += '<div class="file-item"><span class="file-name">' + esc(f.name) + '</span>'
          + '<span class="file-size">' + fmtBytes(f.size) + '</span>'
          + '<button class="btn-icon file-delete" data-file="' + esc(f.id || '') + '" title="Delete">&#128465;</button></div>';
      });
    } else {
      html += '<div style="color:var(--text-dim);font-size:0.75rem;font-style:italic;padding:4px 0">No files in this project</div>';
    }
    el.innerHTML = html;
  } catch (e) {
    el.innerHTML = '<div style="color:var(--crimson-soft);font-size:0.78rem;padding:8px 0">Error: ' + esc(e.message) + '</div>';
  }
}

async function uploadFiles(files) {
  if (!files || !files.length) return;
  if (!S.activeProject) {
    $('upload-status').textContent = 'Create or select a project before uploading files.';
    $('upload-status').className = 'upload-status upload-status--err';
    return;
  }
  const statusEl = $('upload-status');
  statusEl.textContent = 'Uploading...';
  statusEl.className = 'upload-status';
  const errors = [];
  let uploaded = 0;
  for (const f of files) {
    try {
      const form = new FormData();
      form.append('file', f, f.name);
      form.append('project_id', S.activeProject);
      const r = await fetch('/upload', { method: 'POST', body: form });
      const text = await r.text();
      let data;
      try { data = JSON.parse(text); } catch { throw new Error(text); }
      if (!r.ok) throw new Error(data.detail || text);
      uploaded++;
    } catch (e) {
      errors.push(f.name + ': ' + e.message);
    }
  }
  if (errors.length === 0) {
    statusEl.textContent = 'Upload complete (' + uploaded + ' file' + (uploaded !== 1 ? 's' : '') + ').';
    statusEl.className = 'upload-status upload-status--ok';
  } else if (uploaded > 0) {
    statusEl.textContent = uploaded + ' uploaded, ' + errors.length + ' failed: ' + errors.join('; ');
    statusEl.className = 'upload-status upload-status--err';
  } else {
    statusEl.textContent = 'Upload errors: ' + errors.join('; ');
    statusEl.className = 'upload-status upload-status--err';
  }
  await loadProjects();
  await loadFiles();
}

async function ingestDocs() {
  if (!S.activeProject) {
    $('chat-status').textContent = 'Select a project before re-ingesting.';
    return;
  }
  $('chat-status').textContent = 'Ingesting documents...';
  try {
    const res = await POST('/ingest', { project_id: S.activeProject });
    $('chat-status').textContent = 'Ingested: ' + res.files + ' files, ' + res.chunks + ' chunks';
  } catch (e) {
    $('chat-status').textContent = 'Ingest error: ' + e.message;
  }
}

async function deleteProjectFile(fileId) {
  if (!S.activeProject || !fileId) return;
  if (!confirm('Delete this file from the project and remove its chunks from retrieval?')) return;
  try {
    await DEL('/projects/' + encodeURIComponent(S.activeProject) + '/files/' + encodeURIComponent(fileId));
    await loadProjects();
    await loadFiles();
    $('upload-status').textContent = 'File deleted.';
    $('upload-status').className = 'upload-status upload-status--ok';
  } catch (e) {
    $('upload-status').textContent = 'Delete failed: ' + e.message;
    $('upload-status').className = 'upload-status upload-status--err';
  }
}

function syncQueryModeControls() {
  const mode = $('mode-select').value;
  const webResultsInput = $('web-results-input');
  if (!webResultsInput) return;
  webResultsInput.disabled = mode !== 'research';
  webResultsInput.title = mode === 'research'
    ? ''
    : 'Web results are only used in Research mode.';
}

async function checkHealth() {
  try {
    const h = await GET('/health');
    var provider = h.web_search_provider || 'search';
    var configured = h.web_search_configured ? 'ready' : 'unconfigured';
    $('health-bar').textContent = 'ok \u00b7 ' + (h.collection || 'unknown') + ' \u00b7 ' + provider + ' (' + configured + ')';
  } catch {
    $('health-bar').textContent = 'health check failed';
  }
}

// ---- Event Listeners ----

// Auth tabs
document.querySelectorAll('.auth-tab').forEach(function(tab) {
  tab.addEventListener('click', function() {
    document.querySelectorAll('.auth-tab').forEach(function(t) { t.classList.remove('active'); });
    tab.classList.add('active');
    if (tab.dataset.tab === 'login') {
      $('login-form').classList.remove('hidden');
      $('register-form').classList.add('hidden');
    } else {
      $('login-form').classList.add('hidden');
      $('register-form').classList.remove('hidden');
    }
  });
});

// Login form
$('login-form').addEventListener('submit', async function(e) {
  e.preventDefault();
  const errEl = $('login-error');
  errEl.textContent = '';
  try {
    await login($('login-username').value, $('login-password').value);
  } catch (err) {
    errEl.textContent = err.message;
  }
});

// Register form
$('register-form').addEventListener('submit', async function(e) {
  e.preventDefault();
  const errEl = $('register-error');
  errEl.textContent = '';
  const pw = $('reg-password').value;
  const confirm = $('reg-confirm').value;
  if (pw !== confirm) { errEl.textContent = 'Passwords do not match'; return; }
  try {
    await register($('reg-username').value, pw);
  } catch (err) {
    errEl.textContent = err.message;
  }
});

// Logout
$('logout-btn').addEventListener('click', logout);

// Projects
$('create-project-btn').addEventListener('click', createProject);
$('delete-project-btn').addEventListener('click', deleteActiveProject);
$('project-select').addEventListener('change', async function() {
  await setActiveProject(this.value, true);
});
$('project-name-input').addEventListener('keydown', function(e) {
  if (e.key === 'Enter') {
    e.preventDefault();
    createProject();
  }
});

// New conversation
$('new-conv-btn').addEventListener('click', createConversation);

// Conversation list clicks (delegation)
$('conv-list').addEventListener('click', function(e) {
  const accordionBtn = e.target.closest('[data-toggle-accordion]');
  if (accordionBtn) {
    e.stopPropagation();
    const key = accordionBtn.dataset.toggleAccordion;
    S.convoSections[key] = !S.convoSections[key];
    renderConversations();
    return;
  }
  // Export button
  const exportBtn = e.target.closest('.conv-export');
  if (exportBtn) {
    e.stopPropagation();
    const item = exportBtn.closest('.conv-item');
    if (item) showExportDropdown(item.dataset.conv, exportBtn);
    return;
  }
  // Rename button
  const renameBtn = e.target.closest('.conv-rename');
  if (renameBtn) {
    e.stopPropagation();
    const item = renameBtn.closest('.conv-item');
    if (item) startRename(item.dataset.conv);
    return;
  }
  // Delete button
  const deleteBtn = e.target.closest('.conv-delete');
  if (deleteBtn) {
    e.stopPropagation();
    const item = deleteBtn.closest('.conv-item');
    if (item) deleteConversation(item.dataset.conv);
    return;
  }
  // Select conversation
  const item = e.target.closest('.conv-item');
  if (item) selectConversation(item.dataset.conv);
});

// Send message
$('send-btn').addEventListener('click', sendMessage);

// Textarea: Enter to send, Shift+Enter for newline, auto-resize
$('chat-input').addEventListener('keydown', function(e) {
  if (e.key === 'Enter' && !e.shiftKey) {
    e.preventDefault();
    sendMessage();
  }
});
$('chat-input').addEventListener('input', function() {
  this.style.height = 'auto';
  this.style.height = Math.min(this.scrollHeight, 160) + 'px';
});

// File upload
$('file-input').addEventListener('change', async function() {
  await uploadFiles(this.files);
  this.value = '';
});

// Drop zone
const dz = $('drop-zone');
dz.addEventListener('dragover', function(e) { e.preventDefault(); dz.classList.add('active'); });
dz.addEventListener('dragleave', function() { dz.classList.remove('active'); });
dz.addEventListener('drop', async function(e) {
  e.preventDefault();
  dz.classList.remove('active');
  await uploadFiles(e.dataTransfer.files);
});

// Ingest & refresh
$('ingest-btn').addEventListener('click', ingestDocs);
$('refresh-files-btn').addEventListener('click', loadFiles);
$('mode-select').addEventListener('change', syncQueryModeControls);
$('file-list').addEventListener('click', function(e) {
  const btn = e.target.closest('.file-delete');
  if (btn) deleteProjectFile(btn.dataset.file);
});

// Mobile toggles
$('hamburger-btn').addEventListener('click', function() {
  $('conv-sidebar').classList.toggle('open');
  $('right-panel').classList.remove('open');
});
$('panel-toggle-btn').addEventListener('click', function() {
  $('right-panel').classList.toggle('open');
  $('conv-sidebar').classList.remove('open');
});

// ---- Export ----
function showExportDropdown(convId, anchor) {
  var old = document.querySelector('.export-dropdown');
  if (old) old.remove();
  var dd = document.createElement('div');
  dd.className = 'export-dropdown';
  [['csv','Export CSV'],['txt','Export TXT'],['docx','Export DOCX']].forEach(function(f) {
    var a = document.createElement('a');
    a.href = '#';
    a.textContent = f[1];
    a.addEventListener('click', function(e) {
      e.preventDefault();
      window.open('/export/' + encodeURIComponent(convId) + '?format=' + f[0], '_blank');
      dd.remove();
    });
    dd.appendChild(a);
  });
  var rect = anchor.getBoundingClientRect();
  dd.style.position = 'fixed';
  dd.style.top = (rect.bottom + 4) + 'px';
  dd.style.left = rect.left + 'px';
  document.body.appendChild(dd);
  setTimeout(function() {
    document.addEventListener('click', function close(e) {
      if (!dd.contains(e.target)) { dd.remove(); document.removeEventListener('click', close); }
    });
  }, 0);
}

// ---- Generate File ----
$('generate-btn').addEventListener('click', function() {
  $('generate-overlay').classList.remove('hidden');
  $('gen-prompt').value = '';
  $('gen-prompt').focus();
});
$('gen-cancel').addEventListener('click', function() {
  $('generate-overlay').classList.add('hidden');
});
$('generate-overlay').addEventListener('click', function(e) {
  if (e.target === this) $('generate-overlay').classList.add('hidden');
});
$('gen-submit').addEventListener('click', async function() {
  var prompt = $('gen-prompt').value.trim();
  if (!prompt) return;
  var fmt = $('gen-format').value;
  $('gen-submit').disabled = true;
  $('gen-submit').textContent = 'Generating...';
  try {
    var body = { prompt: prompt, format: fmt };
    if (S.activeConvo) body.convo_id = S.activeConvo;
    var resp = await fetch('/generate-file', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify(body),
    });
    if (!resp.ok) {
      var text = await resp.text();
      var detail;
      try { detail = JSON.parse(text).detail; } catch { detail = text; }
      throw new Error(detail || 'Generation failed');
    }
    var blob = await resp.blob();
    var cd = resp.headers.get('content-disposition') || '';
    var match = cd.match(/filename="?([^"]+)"?/);
    var filename = match ? match[1] : ('generated.' + fmt);
    var url = URL.createObjectURL(blob);
    var a = document.createElement('a');
    a.href = url; a.download = filename;
    document.body.appendChild(a); a.click(); a.remove();
    URL.revokeObjectURL(url);
    $('generate-overlay').classList.add('hidden');
  } catch (e) {
    alert('Generation error: ' + e.message);
  } finally {
    $('gen-submit').disabled = false;
    $('gen-submit').textContent = 'Generate';
  }
});

// ---- Init ----
syncQueryModeControls();
checkSession();

})();
</script>
</body>
</html>"""
